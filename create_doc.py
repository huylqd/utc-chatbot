# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def apply_normal(para, first_indent=True):
    fmt = para.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.2
    if first_indent:
        fmt.first_line_indent = Cm(1.0)
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

def fmt_run(run, bold=False, italic=False, size=13):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_normal(doc, text, first_indent=True, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    fmt_run(run, bold=bold)
    return p

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.bold = True
    return p

def add_heading2(doc, text, numbered=True):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 3']
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    return p

def add_list_item(doc, text, level=1, bold=False):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    indent = Cm(1.0 * level)
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(text)
    fmt_run(run, bold=bold)
    return p

def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'• {text}')
    fmt_run(run, bold=bold)
    return p

def add_caption(doc, text, is_table=True):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.bold = False
    return p

def add_tof_entry(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def make_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, 'BDD7EE')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    return table

# ── MAIN ─────────────────────────────────────────────────────────────────────
doc = Document()

# ── Page setup Section 0 (cover) ─────────────────────────────────────────────
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(1.5)

# ═══════════════════════════════════════════════════════════════════════════════
# TRANG BÌA
# ═══════════════════════════════════════════════════════════════════════════════
def center_bold(doc, text, size=13, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = True
    return p

center_bold(doc, 'HỌC VIỆN KỸ THUẬT MẬT MÃ', 14, 4)
center_bold(doc, 'KHOA CÔNG NGHỆ THÔNG TIN', 14, 4)

p_sep = doc.add_paragraph()
p_sep.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sep.paragraph_format.space_after = Pt(30)
p_sep.paragraph_format.first_line_indent = Cm(0)
run = p_sep.add_run('───────────────────')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)

center_bold(doc, 'ĐỒ ÁN TỐT NGHIỆP', 18, 6)

p_blank = doc.add_paragraph()
p_blank.paragraph_format.space_after = Pt(12)

# Tên đề tài
p_title = doc.add_paragraph()
p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(30)
p_title.paragraph_format.first_line_indent = Cm(0)
run = p_title.add_run('XÂY DỰNG HỆ THỐNG CHATBOT HỖ TRỢ TRUY VẤN ĐIỂM\nCỦA SINH VIÊN SỬ DỤNG KỸ THUẬT RAG')
run.font.name = 'Times New Roman'; run.font.size = Pt(16); run.font.bold = True

# Thông tin sinh viên / GVHD
def info_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(3.0)
    p.paragraph_format.first_line_indent = Cm(0)
    r1 = p.add_run(f'{label:<30}')
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(13)
    r2 = p.add_run(f':  {value}')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(13); r2.font.bold = True

info_line(doc, 'Sinh viên thực hiện', 'Bùi Đức Trọng')
info_line(doc, 'Mã sinh viên', 'CT060241')
info_line(doc, 'Lớp', 'CT6B')
info_line(doc, 'Chuyên ngành', 'Phát triển phần mềm di động')
info_line(doc, 'Người hướng dẫn 1', 'ThS. Mạc Văn Hải – Học viện Kỹ thuật Mật mã')
info_line(doc, 'Người hướng dẫn 2', 'TS. Lê Đức Thuận – Học viện Kỹ thuật Mật mã')

p_blank2 = doc.add_paragraph()
p_blank2.paragraph_format.space_after = Pt(30)

center_bold(doc, 'Hà Nội, tháng 6 năm 2026', 13, 6)

# ── Add new section for main content ──────────────────────────────────────────
from docx.oxml import OxmlElement
new_sect_pr = OxmlElement('w:sectPr')
pgMar = OxmlElement('w:pgMar')
pgMar.set(qn('w:top'),    str(int(Cm(2.5).emu / 914)))
pgMar.set(qn('w:bottom'), str(int(Cm(2.5).emu / 914)))
pgMar.set(qn('w:left'),   str(int(Cm(3.0).emu / 914)))
pgMar.set(qn('w:right'),  str(int(Cm(2.0).emu / 914)))
new_sect_pr.append(pgMar)

# Proper section break via paragraph
p_break = doc.add_paragraph()
p_break.paragraph_format.space_before = Pt(0)
p_break.paragraph_format.space_after = Pt(0)
pPr = p_break._p.get_or_add_pPr()
sectPr2 = OxmlElement('w:sectPr')
pgSz = OxmlElement('w:pgSz')
pgSz.set(qn('w:w'), '12240')
pgSz.set(qn('w:h'), '15840')
sectPr2.append(pgSz)
pgMar2 = OxmlElement('w:pgMar')
# Convert cm to twips: 1 cm = 567 twips
pgMar2.set(qn('w:top'),    '1418')   # 2.5cm
pgMar2.set(qn('w:bottom'), '1418')
pgMar2.set(qn('w:left'),   '1701')   # 3.0cm
pgMar2.set(qn('w:right'),  '1134')   # 2.0cm
sectPr2.append(pgMar2)
pPr.append(sectPr2)

# Update main section margins
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

# ═══════════════════════════════════════════════════════════════════════════════
# LỜI CAM ĐOAN
# ═══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'LỜI CAM ĐOAN')

add_normal(doc, 'Em tên là Bùi Đức Trọng, mã sinh viên CT060241, lớp CT6B, chuyên ngành Phát triển phần mềm di động, ngành Công nghệ thông tin, Học viện Kỹ thuật Mật mã.')
add_normal(doc, 'Em xin cam đoan đồ án tốt nghiệp với đề tài: "Xây dựng hệ thống Chatbot hỗ trợ truy vấn điểm của sinh viên sử dụng kỹ thuật RAG" là công trình nghiên cứu do em thực hiện dưới sự hướng dẫn của ThS. Mạc Văn Hải và TS. Lê Đức Thuận, Học viện Kỹ thuật Mật mã.')
add_normal(doc, 'Em xin cam đoan toàn bộ nội dung được trình bày trong đề tài là kết quả nghiên cứu của em, không sao chép từ bất kỳ nguồn nào khác mà không có trích dẫn. Mọi thông tin, tài liệu tham khảo được sử dụng trong đề tài đều được trích dẫn rõ ràng và tuân thủ các quy định về sở hữu trí tuệ. Em xin hoàn toàn chịu trách nhiệm về nội dung của đề tài này.')

p_sign_date = doc.add_paragraph()
p_sign_date.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sign_date.paragraph_format.space_before = Pt(12)
p_sign_date.paragraph_format.space_after = Pt(6)
p_sign_date.paragraph_format.first_line_indent = Cm(0)
r = p_sign_date.add_run('Hà Nội, Ngày      tháng 06 năm 2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.italic = True

p_sign_title = doc.add_paragraph()
p_sign_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sign_title.paragraph_format.space_after = Pt(48)
p_sign_title.paragraph_format.first_line_indent = Cm(0)
r = p_sign_title.add_run('SINH VIÊN THỰC HIỆN')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.bold = True

p_sign_name = doc.add_paragraph()
p_sign_name.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sign_name.paragraph_format.space_after = Pt(6)
p_sign_name.paragraph_format.first_line_indent = Cm(0)
r = p_sign_name.add_run('Bùi Đức Trọng')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.font.bold = True

# ═══════════════════════════════════════════════════════════════════════════════
# LỜI NÓI ĐẦU
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'LỜI NÓI ĐẦU')

add_normal(doc, 'Trong bối cảnh chuyển đổi số mạnh mẽ hiện nay, giáo dục đại học đang chứng kiến sự mở rộng nhanh chóng về quy mô đào tạo cùng sự gia tăng khổng lồ của lượng thông tin hành chính và học vụ cần xử lý. Sinh viên ngày càng có nhu cầu tra cứu thông tin điểm số, quy chế đào tạo, thủ tục hành chính một cách nhanh chóng, chính xác và không phụ thuộc vào giờ làm việc của các phòng ban.')
add_normal(doc, 'Phương thức hỗ trợ truyền thống thông qua cán bộ nhân viên hoặc tra cứu thủ công trên các hệ thống phân tán không còn đáp ứng đủ nhu cầu, thường gây lãng phí thời gian và tạo áp lực công việc cho các bộ phận chức năng. Việc xây dựng hệ thống Chatbot hỗ trợ thông minh không chỉ giúp tối ưu hóa trải nghiệm của người dùng mà còn giảm tải đáng kể áp lực cho hệ thống quản lý hành chính của nhà trường.')
add_normal(doc, 'Đề tài "Xây dựng hệ thống Chatbot hỗ trợ truy vấn điểm của sinh viên sử dụng kỹ thuật RAG" tập trung nghiên cứu và xây dựng một trợ lý ảo phục vụ cộng đồng sinh viên và giảng viên Học viện Kỹ thuật Mật mã, có khả năng tra cứu điểm số theo thời gian thực và giải đáp các vấn đề học vụ, quy chế đào tạo dựa trên kho tài liệu nội bộ.')
add_normal(doc, 'Đồ án được cấu trúc thành bốn chương chính:')
add_bullet(doc, 'Chương 1: Tổng quan đề tài và cơ sở lý thuyết – Khảo sát thực trạng các hệ thống Chatbot giáo dục, tìm hiểu về LLM/SLM, kỹ thuật RAG, Graph RAG và kiến trúc Agent.')
add_bullet(doc, 'Chương 2: Phương pháp và kỹ thuật đề xuất – Mô tả chi tiết kiến trúc hệ thống đa tác tử, chiến lược xây dựng Graph RAG, tiền xử lý dữ liệu và tích hợp LLM.')
add_bullet(doc, 'Chương 3: Phân tích, thiết kế và cài đặt hệ thống – Phân tích yêu cầu, thiết kế use case, kiến trúc tổng thể, cơ sở dữ liệu, API và giao diện người dùng.')
add_bullet(doc, 'Chương 4: Kết quả và thực nghiệm – Môi trường thử nghiệm, đánh giá độ chính xác truy xuất, so sánh RAG vs GraphRAG và demo chức năng hệ thống.')
add_normal(doc, 'Việc ứng dụng AI và Graph RAG để quản lý tri thức và hỗ trợ truy vấn thông tin là một xu thế tất yếu, giúp đẩy nhanh quá trình chuyển đổi số trong giáo dục đại học. Em xin chân thành cảm ơn ThS. Mạc Văn Hải và TS. Lê Đức Thuận đã tận tình hướng dẫn, cũng như toàn thể thầy cô và bạn bè đã hỗ trợ trong suốt quá trình thực hiện đồ án.')

# ═══════════════════════════════════════════════════════════════════════════════
# MỤC LỤC
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'MỤC LỤC')
toc_entries = [
    ('LỜI CAM ĐOAN', ''),
    ('LỜI NÓI ĐẦU', ''),
    ('MỤC LỤC', ''),
    ('DANH MỤC TỪ VIẾT TẮT', ''),
    ('DANH MỤC BẢNG BIỂU', ''),
    ('DANH MỤC HÌNH ẢNH', ''),
    ('CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI VÀ CƠ SỞ LÝ THUYẾT', '1'),
    ('  1.1. Khảo sát hệ thống Chatbot hiện có', '1'),
    ('  1.2. Tìm hiểu về Large Language Models và Small Language Models', '9'),
    ('  1.3. Kỹ thuật Prompt Engineering', '16'),
    ('  1.4. Tìm hiểu về Retrieval-Augmented Generation (RAG)', '21'),
    ('  1.5. Tìm hiểu về Knowledge Graph', '35'),
    ('  1.6. Tìm hiểu về Agent và Multi-Agent Systems', '39'),
    ('  1.7. Tổng kết chương 1', '42'),
    ('CHƯƠNG 2. PHƯƠNG PHÁP VÀ KỸ THUẬT ĐỀ XUẤT', '43'),
    ('  2.1. Tổng quan phương pháp đề xuất', '43'),
    ('  2.2. Thu thập và tiền xử lý dữ liệu', '46'),
    ('  2.3. Xây dựng hệ thống Graph RAG', '48'),
    ('  2.4. Tổng kết chương 2', '53'),
    ('CHƯƠNG 3. PHÂN TÍCH, THIẾT KẾ VÀ CÀI ĐẶT HỆ THỐNG', '55'),
    ('  3.1. Phân tích hệ thống', '55'),
    ('  3.2. Thiết kế hệ thống', '69'),
    ('  3.3. Xây dựng và cài đặt phần mềm', '85'),
    ('  3.4. Tổng kết chương 3', '89'),
    ('CHƯƠNG 4. KẾT QUẢ VÀ THỰC NGHIỆM', '97'),
    ('  4.1. Môi trường thực nghiệm', '97'),
    ('  4.2. Đánh giá mô hình Retrieval', '99'),
    ('  4.3. Đánh giá hiệu năng hệ thống', '103'),
    ('  4.4. Demo chức năng hệ thống', '104'),
    ('  4.5. Phân tích và thảo luận', '108'),
    ('  4.6. Tổng kết chương 4', '110'),
    ('KẾT LUẬN', '111'),
    ('TÀI LIỆU THAM KHẢO', '114'),
    ('PHỤ LỤC', '117'),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    is_chapter = not entry.startswith('  ')
    lvl_indent = Cm(0) if is_chapter else Cm(0.5)
    p.paragraph_format.left_indent = lvl_indent
    run = p.add_run(entry.strip())
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = is_chapter

# ═══════════════════════════════════════════════════════════════════════════════
# DANH MỤC TỪ VIẾT TẮT
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'DANH MỤC TỪ VIẾT TẮT')

abbr_data = [
    ('AI', 'Artificial Intelligence – Trí tuệ nhân tạo'),
    ('API', 'Application Programming Interface – Giao diện lập trình ứng dụng'),
    ('BFS', 'Breadth-First Search – Thuật toán tìm kiếm theo chiều rộng'),
    ('CPU', 'Central Processing Unit – Bộ xử lý trung tâm'),
    ('FAISS', 'Facebook AI Similarity Search – Thư viện tìm kiếm tương đồng vector'),
    ('GA', 'Generally Available – Phiên bản phát hành chính thức'),
    ('GPT', 'Generative Pre-trained Transformer'),
    ('GPU', 'Graphics Processing Unit – Đơn vị xử lý đồ họa'),
    ('Graph RAG', 'Graph Retrieval-Augmented Generation'),
    ('JWT', 'JSON Web Token – Chuẩn xác thực dựa trên token'),
    ('KG', 'Knowledge Graph – Đồ thị tri thức'),
    ('LLM', 'Large Language Model – Mô hình ngôn ngữ lớn'),
    ('LoRA', 'Low-Rank Adaptation – Kỹ thuật tinh chỉnh mô hình hiệu quả'),
    ('LSTM', 'Long Short-Term Memory – Mạng nhớ dài ngắn hạn'),
    ('MAS', 'Multi-Agent System – Hệ thống đa tác tử'),
    ('ML', 'Machine Learning – Học máy'),
    ('MoE', 'Mixture of Experts – Kiến trúc hỗn hợp chuyên gia'),
    ('NLP', 'Natural Language Processing – Xử lý ngôn ngữ tự nhiên'),
    ('OCR', 'Optical Character Recognition – Nhận dạng ký tự quang học'),
    ('PDF', 'Portable Document Format'),
    ('QLoRA', 'Quantized Low-Rank Adaptation'),
    ('RAG', 'Retrieval-Augmented Generation – Sinh tăng cường truy xuất'),
    ('ReAct', 'Reasoning and Acting – Lý luận và hành động'),
    ('REST', 'Representational State Transfer – Kiến trúc API phổ biến'),
    ('RNN', 'Recurrent Neural Network – Mạng nơ-ron hồi quy'),
    ('SLM', 'Small Language Model – Mô hình ngôn ngữ nhỏ'),
    ('VRAM', 'Video Random Access Memory – Bộ nhớ video'),
]

tbl = doc.add_table(rows=1+len(abbr_data), cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hrow = tbl.rows[0]
for i, h in enumerate(['Từ viết tắt', 'Giải nghĩa']):
    set_cell_bg(hrow.cells[i], 'BDD7EE')
    p = hrow.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h); run.font.name='Times New Roman'; run.font.size=Pt(12); run.font.bold=True
for ri, (abbr, full) in enumerate(abbr_data):
    row = tbl.rows[ri+1]
    p0 = row.cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(abbr); r0.font.name='Times New Roman'; r0.font.size=Pt(11); r0.font.bold=True
    p1 = row.cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(full); r1.font.name='Times New Roman'; r1.font.size=Pt(11)
for row in tbl.rows:
    row.cells[0].width = Cm(3.0)
    row.cells[1].width = Cm(12.5)

# ═══════════════════════════════════════════════════════════════════════════════
# DANH MỤC BẢNG BIỂU
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'DANH MỤC BẢNG BIỂU')

bang_entries = [
    'Bảng 1.1. So sánh các hệ thống Chatbot giáo dục tiêu biểu',
    'Bảng 1.2. So sánh các hướng tiếp cận xây dựng hệ thống Chatbot hỏi đáp tài liệu',
    'Bảng 1.3. So sánh toàn diện LLM và SLM',
    'Bảng 1.4. So sánh chi tiết các mô hình LLM/SLM tiêu biểu (cập nhật Q1/2026)',
    'Bảng 1.5. Bảng thành phần cấu tạo của một prompt',
    'Bảng 3.1. Bảng đặc tả usecase đăng ký',
    'Bảng 3.2. Bảng đặc tả usecase đăng nhập',
    'Bảng 3.3. Bảng đặc tả usecase chat với AI',
    'Bảng 3.4. Bảng đặc tả usecase quản lý lịch sử hội thoại',
    'Bảng 3.5. Bảng đặc tả usecase quản lý file cá nhân',
    'Bảng 3.6. Bảng đặc tả usecase xem danh sách mô hình LLM',
    'Bảng 3.7. Bảng đặc tả usecase quản lý mô hình LLM',
    'Bảng 3.8. Bảng đặc tả usecase quản lý dữ liệu RAG',
    'Bảng 3.9. Bảng đặc tả usecase quản lý người dùng',
    'Bảng 3.10. Bảng đặc tả usecase giới hạn hệ thống',
    'Bảng 3.11. Bảng đặc tả usecase thống kê sử dụng',
    'Bảng 3.12. Cấu trúc bảng Score – lưu trữ thông tin điểm số',
    'Bảng 3.13. Cấu trúc bảng Structure – lưu trữ thông tin môn học',
    'Bảng 3.14. Cấu trúc bảng Student – lưu trữ thông tin sinh viên',
    'Bảng 3.15. Cấu trúc bảng Users – lưu trữ tài khoản và xác thực',
    'Bảng 3.16. Cấu trúc bảng Conversations – lưu trữ cuộc trò chuyện',
    'Bảng 3.17. Cấu trúc bảng Messages – lưu trữ tin nhắn',
    'Bảng 3.18. Cấu trúc bảng Rate_limits – kiểm soát tốc độ yêu cầu',
    'Bảng 3.19. Cấu trúc bảng LLM_models – quản lý mô hình AI',
    'Bảng 3.20. Cấu trúc bảng User_document_embeddings',
    'Bảng 3.21. API tra cứu điểm sinh viên',
    'Bảng 3.22. API tạo cuộc trò chuyện mới',
    'Bảng 3.23. API xoá cuộc trò chuyện',
    'Bảng 3.24. API lấy lịch sử chat của cuộc trò chuyện',
    'Bảng 3.25. API gửi tin nhắn trong cuộc trò chuyện',
    'Bảng 3.26. API đăng nhập hệ thống',
    'Bảng 3.27. API lấy thông tin người dùng hiện tại',
    'Bảng 3.28. API làm mới token',
    'Bảng 3.29. API tạo tài khoản người dùng',
    'Bảng 3.30. API xem danh sách người dùng',
    'Bảng 3.31. API cập nhật người dùng',
    'Bảng 3.32. API xoá người dùng',
    'Bảng 3.33. API xem model đang hoạt động',
    'Bảng 3.34. API kích hoạt model',
    'Bảng 3.35. API tải lên model mới',
    'Bảng 3.36. API lấy cấu hình giới hạn tốc độ',
    'Bảng 3.37. API tạo/cập nhật cấu hình giới hạn tốc độ',
    'Bảng 3.38. API xem thống kê rate limit tất cả người dùng',
    'Bảng 3.39. API xem tổng hợp sử dụng rate limit',
    'Bảng 3.40. API xem thống kê dashboard',
    'Bảng 3.41. API xem thống kê cuộc hội thoại',
    'Bảng 3.42. API tải file huấn luyện RAG',
    'Bảng 3.43. API liệt kê file huấn luyện RAG',
    'Bảng 3.44. API xoá file huấn luyện RAG',
    'Bảng 3.45. API rebuild toàn bộ RAG index',
    'Bảng 3.46. API rebuild RAG index theo phòng ban',
    'Bảng 3.47. API liệt kê phòng ban dữ liệu RAG',
    'Bảng 3.48. API tạo thư mục dữ liệu RAG',
    'Bảng 3.49. API xoá thư mục dữ liệu RAG',
    'Bảng 3.50. API tải file từ thư mục dữ liệu RAG',
    'Bảng 4.1. Chi tiết bộ dữ liệu hành chính thử nghiệm',
    'Bảng 4.2. Mô tả bộ dữ liệu đánh giá',
    'Bảng 4.3. Kết quả thực nghiệm kiến trúc RAG kết hợp LLMs',
    'Bảng 4.4. Kết quả thực nghiệm kiến trúc GraphRAG kết hợp LLMs',
]
for entry in bang_entries:
    add_tof_entry(doc, entry)

# ═══════════════════════════════════════════════════════════════════════════════
# DANH MỤC HÌNH ẢNH
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'DANH MỤC HÌNH ẢNH')

hinh_entries = [
    'Hình 1.1. Kiến trúc tổng quan hệ thống RAG',
    'Hình 1.2. Pipeline xử lý tổng thể của Graph RAG',
    'Hình 1.3. Minh họa Knowledge Graph kết nối thực thể qua các quan hệ ngữ nghĩa',
    'Hình 2.1. Kiến trúc xử lý tổng thể hệ thống đa tác tử',
    'Hình 2.2. Mẫu đồ thị tri thức cho tập dữ liệu có kiến trúc phức tạp',
    'Hình 2.3. Đồ thị tri thức sau khi được phân cụm cộng đồng',
    'Hình 2.4. Sơ đồ luồng truy vấn Graph RAG',
    'Hình 2.5. Luồng hoạt động của ReAct Agent',
    'Hình 3.1. Biểu đồ use case tổng quát hệ thống',
    'Hình 3.2. Biểu đồ hoạt động chức năng đăng ký',
    'Hình 3.3. Biểu đồ hoạt động chức năng đăng nhập',
    'Hình 3.4. Biểu đồ hoạt động chức năng chat với AI',
    'Hình 3.5. Biểu đồ hoạt động chức năng quản lý lịch sử hội thoại',
    'Hình 3.6. Biểu đồ hoạt động chức năng quản lý file cá nhân',
    'Hình 3.7. Biểu đồ hoạt động chức năng xem danh sách mô hình',
    'Hình 3.8. Biểu đồ hoạt động chức năng quản lý mô hình LLM',
    'Hình 3.9. Biểu đồ hoạt động chức năng quản lý dữ liệu RAG',
    'Hình 3.10. Biểu đồ hoạt động chức năng quản lý người dùng',
    'Hình 3.11. Biểu đồ hoạt động chức năng giới hạn hệ thống',
    'Hình 3.12. Biểu đồ hoạt động chức năng thống kê sử dụng',
    'Hình 3.13. Kiến trúc tổng thể hệ thống',
    'Hình 3.14. Kiến trúc Graph RAG',
    'Hình 3.15. Giao diện trang chủ hệ thống',
    'Hình 3.16. Giao diện hội thoại chính',
    'Hình 3.17. Giao diện tra cứu điểm',
    'Hình 3.18. Giao diện truy vấn tài liệu hành chính nội bộ',
    'Hình 3.19. Giao diện hỏi đáp với tệp tin tải lên',
    'Hình 3.20. Giao diện quản lý tệp tin tải lên',
    'Hình 3.21. Giao diện trang chủ Admin',
    'Hình 3.22. Giao diện quản lý người dùng',
    'Hình 3.23. Giao diện quản lý lịch sử hội thoại',
    'Hình 3.24. Giao diện giới hạn sử dụng',
    'Hình 3.25. Giao diện quản lý mô hình LLM',
    'Hình 3.26. Giao diện giới hạn tệp tin tải lên',
    'Hình 3.27. Giao diện quản lý dữ liệu RAG',
    'Hình 4.1. Giao diện khung chat và thao tác gửi câu hỏi',
    'Hình 4.2. Kết quả phản hồi của Chatbot bám sát tài liệu nội bộ',
    'Hình 4.3. Giao diện danh sách lịch sử hội thoại của người dùng',
    'Hình 4.4. Giao diện tải file và quản lý danh sách tài liệu tri thức',
    'Hình 4.5. Dashboard thống kê hoạt động của Chatbot',
]
for entry in hinh_entries:
    add_tof_entry(doc, entry)

print("Phase 1 done – cover, cam doan, loi noi dau, muc luc, danh muc added.")
doc.save('document/TAI_LIEU_SAN_PHAM_CHUAN_FORMAT_part1.docx')
print("Saved part1.")
