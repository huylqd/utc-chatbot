import glob
import os
import io
import sys
import logging
from typing import List, Optional, Dict, Any
import tempfile
import re

from langchain_community.retrievers import BM25Retriever
from langchain_core.documents import Document
from langchain_core.retrievers import BaseRetriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from pydantic import Field, BaseModel
from llm.config import get_llm  # Import get_llm thay vì get_gemini_llm

# Set up logging
logger = logging.getLogger(__name__)

# Import metadata configuration
from .metadata_config import get_metadata_config
from .semantic_analyzer import analyze_query_semantic_filter
from .docling_extractor import extract_text_with_docling, is_docling_available
from dotenv import load_dotenv
load_dotenv()
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
# Optional imports for file processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Set UTF-8 encoding
# sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
# sys.stdin = io.TextIOWrapper(sys.stdin.buffer, encoding='utf-8')


class MetadataEnhancedHybridRetriever(BaseRetriever, BaseModel):
    vectorstore: FAISS = Field(description="FAISS vector store")
    bm25_retriever: BM25Retriever = Field(description="BM25 retriever")
    k: int = Field(default=4, description="Number of documents to retrieve")
    window_size: int = Field(default=1, description="Number of adjacent chunks to include (sliding window)")
    all_documents: Optional[List[Document]] = Field(default=None, description="All documents for sliding window")
    
    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(self, query: str, metadata_filter: Optional[Dict[str, Any]] = None) -> List[Document]:
        """Retrieve documents with optional metadata filtering and sliding window"""
        
        # Vector search with metadata filtering if supported
        if metadata_filter:
            try:
                vector_docs = self.vectorstore.similarity_search(
                    query, 
                    k=self.k,
                    filter=metadata_filter
                )
            except Exception:
                # Fallback if filtering not supported
                vector_docs = self.vectorstore.similarity_search(query, k=self.k)
                # Filter afterward
                filtered_vector_docs = []
                for doc in vector_docs:
                    match = True
                    for key, value in metadata_filter.items():
                        if key not in doc.metadata or doc.metadata[key] != value:
                            match = False
                            break
                    if match:
                        filtered_vector_docs.append(doc)
                vector_docs = filtered_vector_docs
        else:
            vector_docs = self.vectorstore.similarity_search(query, k=self.k)
        
        # BM25 search (filter afterward since BM25Retriever doesn't support metadata filtering)
        bm25_docs = self.bm25_retriever.invoke(query)
        
        # Filter BM25 results by metadata if specified
        if metadata_filter and hasattr(self.bm25_retriever, 'docs'):
            filtered_bm25_docs = []
            for doc in bm25_docs:
                match = True
                for key, value in metadata_filter.items():
                    if key not in doc.metadata or doc.metadata[key] != value:
                        match = False
                        break
                if match:
                    filtered_bm25_docs.append(doc)
            bm25_docs = filtered_bm25_docs

        # Get initial relevant documents
        initial_docs = vector_docs + bm25_docs
        
        # Apply sliding window to get adjacent chunks - but preserve ranking
        if self.window_size > 0 and self.all_documents:
            expanded_docs = self._apply_sliding_window_smart(initial_docs)
        else:
            expanded_docs = initial_docs

        # Combine and deduplicate
        all_docs = []
        seen_content = set()
        for doc in expanded_docs:
            if doc.page_content not in seen_content:
                all_docs.append(doc)
                seen_content.add(doc.page_content)

        return all_docs
    
    def _apply_sliding_window_smart(self, initial_docs: List[Document]) -> List[Document]:
        """Apply sliding window smartly - preserve ranking and avoid dilution"""
        expanded_docs = []
        seen_positions = set()
        
        # Priority 1: Keep original documents in their original order
        for doc in initial_docs:
            expanded_docs.append(doc)
            # Mark the position as seen
            current_idx = self._find_document_index(doc)
            if current_idx is not None:
                seen_positions.add(current_idx)
        
        # Priority 2: Add adjacent chunks only for high-ranking documents (top 30%)
        high_priority_count = max(1, len(initial_docs) // 3)
        
        for i, doc in enumerate(initial_docs[:high_priority_count]):
            current_idx = self._find_document_index(doc)
            
            if current_idx is not None:
                # Get window range - use full configured window size for complete coverage
                window_size = self.window_size  # Use full window size from config
                start_idx = max(0, current_idx - window_size)
                end_idx = min(len(self.all_documents), current_idx + window_size + 1)
                
                # Add adjacent documents that aren't already included
                for j in range(start_idx, end_idx):
                    if j not in seen_positions:
                        expanded_docs.append(self.all_documents[j])
                        seen_positions.add(j)
        
        return expanded_docs
    
    def _apply_sliding_window(self, initial_docs: List[Document]) -> List[Document]:
        """Apply sliding window to get adjacent chunks"""
        expanded_docs = []
        
        for doc in initial_docs:
            # Find the index of current document in all_documents
            current_idx = self._find_document_index(doc)
            
            if current_idx is not None:
                # Get window range
                start_idx = max(0, current_idx - self.window_size)
                end_idx = min(len(self.all_documents), current_idx + self.window_size + 1)
                
                # Add all documents in the window
                for i in range(start_idx, end_idx):
                    expanded_docs.append(self.all_documents[i])
            else:
                # If not found in all_documents, add the original document
                expanded_docs.append(doc)
        
        return expanded_docs
    
    def _find_document_index(self, target_doc: Document) -> Optional[int]:
        """Find the index of a document in all_documents by comparing content and metadata"""
        if not self.all_documents:
            return None
            
        for i, doc in enumerate(self.all_documents):
            # Compare by content and source path (more reliable than exact content match)
            if (doc.page_content == target_doc.page_content or
                (hasattr(doc, 'metadata') and hasattr(target_doc, 'metadata') and
                 doc.metadata.get('source_path') == target_doc.metadata.get('source_path') and
                 doc.metadata.get('chunk_index') == target_doc.metadata.get('chunk_index'))):
                return i
        
        return None

    async def _aget_relevant_documents(self, query: str) -> List[Document]:
        raise NotImplementedError("Async retrieval not implemented")


# Keep original class for backward compatibility
class HybridRetriever(BaseRetriever, BaseModel):
    vectorstore: FAISS = Field(description="FAISS vector store")
    bm25_retriever: BM25Retriever = Field(description="BM25 retriever")
    k: int = Field(default=4, description="Number of documents to retrieve")

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(self, query: str) -> List[Document]:
        vector_docs = self.vectorstore.similarity_search(query, k=self.k)
        bm25_docs = self.bm25_retriever.invoke(query)

        all_docs = []
        seen_content = set()
        for doc in vector_docs + bm25_docs:
            if doc.page_content not in seen_content:
                all_docs.append(doc.page_content)
                seen_content.add(doc.page_content)

        return [Document(page_content=content) for content in all_docs]

    async def _aget_relevant_documents(self, query: str) -> List[Document]:
        raise NotImplementedError("Async retrieval not implemented")


def extract_metadata_from_path(file_path: str, base_data_dir: str) -> Dict[str, str]:
    """Extract metadata from file path structure using dynamic configuration"""
    config = get_metadata_config()
    metadata = {}
    
    # Get relative path from base data directory
    rel_path = os.path.relpath(file_path, base_data_dir)
    path_parts = rel_path.split(os.sep)
    
    # Extract filename and extension
    filename = os.path.basename(file_path)
    metadata['filename'] = filename
    metadata['file_extension'] = os.path.splitext(filename)[1]
    
    # Start with default metadata
    metadata.update(config.get_default_metadata())
    
    # Check if first part matches any configured folder
    if len(path_parts) >= 2:
        folder_name = path_parts[0]
        folder_mapping = config.get_folder_mapping(folder_name)
        
        if folder_mapping:
            # Apply folder-level metadata
            for key, value in folder_mapping.items():
                if key != 'subfolders':
                    metadata[key] = value
            
            # Check for subfolder mappings
            if len(path_parts) >= 3 and 'subfolders' in folder_mapping:
                subfolder_name = path_parts[1]
                subfolder_mapping = folder_mapping['subfolders'].get(subfolder_name, {})
                
                if subfolder_mapping:
                    # Apply subfolder-level metadata
                    metadata.update(subfolder_mapping)
                else:
                    # If subfolder not configured, use folder name as level
                    metadata['custom_level'] = subfolder_name
                    metadata['custom_level_path'] = f"{folder_name}/{subfolder_name}"
        else:
            # If folder not configured, use dynamic metadata
            metadata['department'] = folder_name
            metadata['department_vn'] = folder_name.title()
            metadata['source_type'] = 'custom'
            
            # If there's a subfolder, treat it as a level
            if len(path_parts) >= 3:
                subfolder_name = path_parts[1]
                metadata['custom_level'] = subfolder_name
                metadata['custom_level_vn'] = subfolder_name.title()
    
    # Add full path for reference
    metadata['source_path'] = rel_path
    metadata['folder_depth'] = len(path_parts) - 1  # Exclude filename
    
    return metadata


def analyze_query_for_metadata_filter(query: str) -> Dict[str, Any]:
    """
    DEPRECATED: Use semantic analysis instead of keyword matching
    Kept for backward compatibility - now delegates to semantic analyzer
    """
    print("⚠️  Using legacy keyword matching - consider upgrading to semantic analysis")
    return analyze_query_semantic_filter(query, confidence_threshold=0.65)


def analyze_query_for_metadata_filter_legacy(query: str) -> Dict[str, Any]:
    """Legacy keyword-based analysis - kept as fallback only"""
    config = get_metadata_config()
    query_keywords = config.get_query_keywords()
    filters = {}
    query_lower = query.lower()
    
    # Check education level keywords
    if 'education_levels' in query_keywords:
        for level, keywords in query_keywords['education_levels'].items():
            if any(keyword in query_lower for keyword in keywords):
                filters['education_level'] = level
                break
    
    # Check department keywords  
    if 'departments' in query_keywords:
        for dept, keywords in query_keywords['departments'].items():
            if any(keyword in query_lower for keyword in keywords):
                filters['department'] = dept
                break
    
    # Check for custom levels and departments
    # This allows for dynamic detection of new folder structures
    folder_mappings = config.config.get('folder_mappings', {})
    for folder_name, folder_config in folder_mappings.items():
        # Check department Vietnamese name
        dept_vn = folder_config.get('department_vn', '')
        if dept_vn and dept_vn.lower() in query_lower:
            filters['department'] = folder_name
        
        # Check subfolder levels
        if 'subfolders' in folder_config:
            for subfolder, subfolder_config in folder_config['subfolders'].items():
                level_vn = subfolder_config.get('education_level_vn', '')
                if level_vn and level_vn.lower() in query_lower:
                    filters['education_level'] = subfolder
    
    return filters


def read_file_with_metadata(file_path: str, base_data_dir: str) -> tuple[str, Dict[str, str]]:
    """Read file content and extract metadata using Docling for all file types"""
    try:
        # Extract metadata from path
        metadata = extract_metadata_from_path(file_path, base_data_dir)
        
        # Use Docling to extract text from all file types
        content = ""
        
        if is_docling_available():
            logger.info(f"Using Docling to extract text from {file_path}")
            content = extract_text_with_docling(file_path)
            
            if not content:
                logger.warning(f"Docling failed to extract text from {file_path}")
                # Fallback to reading as plain text for .txt files
                if file_path.endswith('.txt') or file_path.endswith('.md') or file_path.endswith('.markdown'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
        else:
            logger.warning("Docling not available, falling back to legacy text extraction")
            # Fallback to reading as plain text for .txt files only
            if file_path.endswith('.txt') or file_path.endswith('.md') or file_path.endswith('.markdown'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                logger.error(f"Cannot process {file_path} without Docling")
                return "", metadata
        
        if not content:
            logger.warning(f"No content extracted from {file_path}")
            return "", metadata
            
        # Create context header dynamically
        header_parts = []
        
        # Add department info
        dept_vn = metadata.get('department_vn', metadata.get('department', ''))
        if dept_vn and dept_vn != 'general':
            header_parts.append(f"từ {dept_vn}")
        
        # Add level info  
        level_vn = metadata.get('education_level_vn', metadata.get('custom_level_vn', ''))
        if level_vn:
            header_parts.append(level_vn)
        
        # Add custom level if exists
        custom_level = metadata.get('custom_level', '')
        if custom_level and not level_vn:
            header_parts.append(custom_level)
        
        if header_parts:
            context_header = f"[Tài liệu {' - '.join(header_parts)}]\n"
        else:
            context_header = f"[Tài liệu {metadata.get('source_path', '')}]\n"
            
        content = context_header + content
        
        return content, metadata
        
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return "", {}


def read_all_files_with_metadata(data_dir: str) -> List[tuple[str, Dict[str, str]]]:
    """Read all files and return content with metadata"""
    file_contents = []
    
    # Process all .txt, .md, .markdown, .pdf, and .docx files recursively
    for root, dirs, files in os.walk(data_dir):
        for file in files:
            if file.endswith(('.txt', '.md', '.markdown', '.docx')):
                file_path = os.path.join(root, file)
                content, metadata = read_file_with_metadata(file_path, data_dir)
                
                if content:  # Only add if content was successfully read
                    file_contents.append((content, metadata))
    
    return file_contents


def read_all_text_files(data_dir):
    """Đọc toàn bộ nội dung các file .txt và .md trong thư mục và tất cả thư mục con"""
    combined_text = ""
    
    # Đọc file trong thư mục chính (.txt và .md)
    for pattern in ["*.txt", "*.md", "*.markdown"]:
        for file_path in glob.glob(os.path.join(data_dir, pattern)):
            with open(file_path, "r", encoding="utf-8") as f:
                combined_text += f.read() + "\n\n"
    
    # Đọc file từ tất cả các thư mục con
    for root, dirs, files in os.walk(data_dir):
        # Bỏ qua thư mục gốc vì đã xử lý ở trên
        if root == data_dir:
            continue
            
        for file in files:
            if file.endswith(('.txt', '.md', '.markdown')):
                file_path = os.path.join(root, file)
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()
                        # Thêm thông tin về nguồn của nội dung
                        folder_name = os.path.basename(root)
                        combined_text += f"[Từ thư mục: {folder_name}]\n{content}\n\n"
                except Exception as e:
                    print(f"Error reading file {file_path}: {e}")
                    
    return combined_text


def smart_text_chunking(content: str, metadata: Dict[str, str], chunk_settings: Dict[str, Any]) -> List[str]:
    """Enhanced smart text chunking that preserves structure and ensures proper overlap"""
    
    # First, detect and preserve structured content (lists, sections, tables)
    preserved_chunks = detect_and_preserve_structured_content(content, chunk_settings)
    
    if preserved_chunks:
        print(f"📋 Preserved {len(preserved_chunks)} structured sections")
        return preserved_chunks
    
    # Check if content contains tables
    if "[BẢNG DỮ LIỆU]" in content and "[KẾT THÚC BẢNG]" in content:
        return handle_table_content(content, chunk_settings)
    
    # Enhanced normal chunking with improved overlap
    return enhanced_text_chunking(content, chunk_settings)


def create_enhanced_vector_database(output_path: str, data_dir: str = "./data"):
    """Create vector database with metadata support using dynamic configuration"""
    try:
        config = get_metadata_config()
        chunk_settings = config.get_chunk_settings()
        
        # Read all files with metadata
        file_contents = read_all_files_with_metadata(data_dir)
        
        if not file_contents:
            raise ValueError("No files found to process")
        
        # Prepare text splitter with config parameters
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_settings.get('chunk_size', 500),
            chunk_overlap=chunk_settings.get('chunk_overlap', 100),
            length_function=len,
            separators=chunk_settings.get('separators', ["\n\n", "\n", ". ", " ", ""]),
            keep_separator=chunk_settings.get('keep_separator', True)
        )
        
        # Process each file and create documents with metadata
        all_documents = []
        for content, metadata in file_contents:
            # Use smart chunking that preserves table structure
            chunks = smart_text_chunking(content, metadata, chunk_settings)
            
            for i, chunk in enumerate(chunks):
                # Create document with metadata
                chunk_metadata = metadata.copy()
                chunk_metadata['chunk_index'] = i
                chunk_metadata['total_chunks'] = len(chunks)
                
                # Add table indicator if chunk contains table
                if "[BẢNG DỮ LIỆU]" in chunk and "[KẾT THÚC BẢNG]" in chunk:
                    chunk_metadata['contains_table'] = True
                    chunk_metadata['chunk_type'] = 'table'
                else:
                    chunk_metadata['contains_table'] = False
                    chunk_metadata['chunk_type'] = 'text'
                
                doc = Document(
                    page_content=chunk,
                    metadata=chunk_metadata
                )
                all_documents.append(doc)
        
        # Create embeddings
        embeddings = OllamaEmbeddings(model="nomic-embed-text")
        
        # Create FAISS vector store from documents (preserves metadata)
        vectorstore = FAISS.from_documents(all_documents, embeddings)
        
        # Ensure output directory exists
        if os.path.dirname(output_path):
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        else:
            os.makedirs(output_path, exist_ok=True)
        
        # Save vector store
        vectorstore.save_local(output_path)
        
        print(f"Created enhanced vector database with {len(all_documents)} documents")
        return all_documents
        
    except Exception as e:
        print(f"Error creating enhanced vector database: {e}")
        raise


def load_enhanced_vector_database(output_path: str, data_dir: str = "./data"):
    """Load or create enhanced vector database"""
    try:
        embeddings = OllamaEmbeddings(model="nomic-embed-text")
        
        if not os.path.exists(output_path):
            print("Vector database not found, creating new one...")
            documents = create_enhanced_vector_database(output_path, data_dir)
        else:
            print("Loading existing vector database...")
            # Load existing vectorstore
            vectorstore = FAISS.load_local(output_path, embeddings, allow_dangerous_deserialization=True)
            
            # Recreate documents for BM25 (we need the text content)
            file_contents = read_all_files_with_metadata(data_dir)
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=500,
                chunk_overlap=100,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""],
                keep_separator=True
            )
            
            documents = []
            for content, metadata in file_contents:
                chunks = text_splitter.split_text(content)
                for i, chunk in enumerate(chunks):
                    chunk_metadata = metadata.copy()
                    chunk_metadata['chunk_index'] = i
                    chunk_metadata['total_chunks'] = len(chunks)
                    doc = Document(page_content=chunk, metadata=chunk_metadata)
                    documents.append(doc)
            
            return vectorstore, documents
        
        # If we created new database, load it
        vectorstore = FAISS.load_local(output_path, embeddings, allow_dangerous_deserialization=True)
        return vectorstore, documents
        
    except Exception as e:
        print(f"Error loading enhanced vector database: {e}")
        raise


def create_enhanced_hybrid_retriever(vector_db_path: str, data_dir: str = "./data", window_size: int = 1):
    """Create enhanced hybrid retriever with metadata support and sliding window"""
    vectorstore, documents = load_enhanced_vector_database(vector_db_path, data_dir)
    
    # Extract text content for BM25
    texts = [doc.page_content for doc in documents]
    bm25_retriever = BM25Retriever.from_texts(texts=texts, k=15)
    
    # Store documents in BM25 retriever for metadata filtering
    bm25_retriever.docs = documents
    
    return MetadataEnhancedHybridRetriever(
        vectorstore=vectorstore,
        bm25_retriever=bm25_retriever,
        k=15,
        window_size=window_size,
        all_documents=documents
    ), documents


def smart_retrieve(retriever: MetadataEnhancedHybridRetriever, query: str, use_smart_filtering: bool = True) -> List[Document]:
    """Smart retrieval with automatic semantic metadata filtering and fallback to full search"""
    if use_smart_filtering:
        # Use semantic analysis instead of keyword matching
        metadata_filter = analyze_query_semantic_filter(query, confidence_threshold=0.65)
        
        if metadata_filter:
            print(f"🎯 Applying semantic metadata filters: {metadata_filter}")
        initial_results = retriever._get_relevant_documents(query, metadata_filter)
    else:
        initial_results = retriever._get_relevant_documents(query)
    
    # Apply context boosting - promote chunks that are part of the same context
    return apply_context_boosting(initial_results, query)


def retrieve_complete_section(retriever: MetadataEnhancedHybridRetriever, query: str, 
                            section_keywords: List[str] = None, additional_keywords: List[str] = None) -> List[Document]:
    """Retrieve complete section by finding all related chunks - generalized approach"""
    
    # First, get initial results
    initial_results = retriever._get_relevant_documents(query)
    
    # If no specific keywords provided, extract from query
    if not section_keywords:
        # Extract meaningful department/section names from query
        query_lower = query.lower()
        section_keywords = []
        
        # Look for common organizational patterns
        org_patterns = ['phòng ', 'ban ', 'viện ', 'trung tâm ', 'khoa ', 'bộ phận ']
        for pattern in org_patterns:
            if pattern in query_lower:
                # Extract the organization name after the pattern
                start_idx = query_lower.find(pattern)
                remaining = query_lower[start_idx:start_idx+50]
                words = remaining.split()[:4]  # Take up to 4 words
                if words:
                    section_keywords.append(' '.join(words))
    
    if not section_keywords:
        # Fallback to regular retrieval
        print("📋 No specific section keywords found, using regular retrieval")
        return apply_context_boosting(initial_results, query)
    
    # Find chunks that contain any of the section keywords
    section_chunks = []
    section_files = set()
    
    # Look for section chunks in initial results and all documents
    for doc in initial_results:
        content_lower = doc.page_content.lower()
        if any(keyword in content_lower for keyword in section_keywords):
            section_chunks.append(doc)
            section_files.add(doc.metadata.get('filename', ''))
    
    # If we found section chunks, get their neighbors for completeness
    if section_chunks and retriever.all_documents:
        chunk_indices = []
        for doc in section_chunks:
            filename = doc.metadata.get('filename', '')
            chunk_index = doc.metadata.get('chunk_index')
            if filename and chunk_index is not None:
                chunk_indices.append((filename, chunk_index))
        
        # Get neighboring chunks for completeness
        complete_chunks = []
        processed_keys = set()  # Use filename:chunk_index as key to avoid duplicates
        
        for filename, chunk_index in chunk_indices:
            # Add the original chunk and neighbors (adaptive window size)
            window_size = 3  # Default window size
            for offset in range(-1, window_size + 1):
                target_index = chunk_index + offset
                key = f"{filename}:{target_index}"
                
                if key not in processed_keys:
                    for doc in retriever.all_documents:
                        if (doc.metadata.get('filename') == filename and 
                            doc.metadata.get('chunk_index') == target_index):
                            complete_chunks.append(doc)
                            processed_keys.add(key)
                            break
        
        # Sort by filename and chunk index for consistency
        complete_chunks.sort(key=lambda x: (x.metadata.get('filename', ''), 
                                          x.metadata.get('chunk_index', 0)))
        
        # Add other high-relevance documents that aren't already included
        processed_keys_simple = {f"{doc.metadata.get('filename')}:{doc.metadata.get('chunk_index')}" 
                               for doc in complete_chunks}
        remaining_docs = []
        for doc in initial_results:
            doc_key = f"{doc.metadata.get('filename')}:{doc.metadata.get('chunk_index')}"
            if doc_key not in processed_keys_simple and len(remaining_docs) < 10:
                remaining_docs.append(doc)
        
        print(f"📋 Section completion: Found {len(complete_chunks)} complete chunks + {len(remaining_docs)} additional docs")
        return complete_chunks + remaining_docs
    
    # Fallback to regular retrieval with context boosting
    print("📋 Section completion: Falling back to context boosting")
    return apply_context_boosting(initial_results, query)


def apply_context_boosting(documents: List[Document], query: str, boost_factor: float = 1.5) -> List[Document]:
    """Apply context boosting to prioritize related chunks from same document/section"""
    if not documents:
        return documents
    
    # Create a mapping of documents by filename and chunk_index
    doc_chunks = {}
    for i, doc in enumerate(documents):
        filename = doc.metadata.get('filename', '')
        chunk_index = doc.metadata.get('chunk_index')
        
        if filename and chunk_index is not None:
            key = f"{filename}:{chunk_index}"
            doc_chunks[key] = {
                'doc': doc, 
                'original_index': i, 
                'boosted_score': 0,
                'relevance_score': 1.0 / (i + 1)  # Higher score for earlier positions
            }
    
    # Find high-scoring chunks and boost their neighbors more aggressively
    high_score_threshold = 0.3  # Top 30% of results get strongest boost
    top_count = max(1, int(len(documents) * high_score_threshold))
    
    for i in range(top_count):
        doc = documents[i]
        filename = doc.metadata.get('filename', '')
        chunk_index = doc.metadata.get('chunk_index')
        
        if filename and chunk_index is not None:
            # Calculate base boost based on position (earlier = much higher boost)
            position_boost = 2.0 / (i + 1)  # Stronger position weighting
            
            # Boost score for current chunk significantly
            current_key = f"{filename}:{chunk_index}"
            if current_key in doc_chunks:
                doc_chunks[current_key]['boosted_score'] += position_boost * 3.0
            
            # Extract topic keywords for better matching
            topic_keywords = extract_topic_keywords(doc.page_content, query)
            
            # Boost neighboring chunks with much stronger weighting
            for offset in [-2, -1, 1, 2]:
                neighbor_index = chunk_index + offset
                neighbor_key = f"{filename}:{neighbor_index}"
                
                if neighbor_key in doc_chunks:
                    neighbor_doc = doc_chunks[neighbor_key]['doc']
                    
                    # Apply strong distance boost
                    distance_boost = boost_factor / abs(offset)
                    
                    # Massive boost for immediate neighbors (consecutive sections)
                    if abs(offset) == 1:
                        distance_boost *= 3.0  # Triple boost for immediate neighbors
                    
                    # Extra boost if neighbor has similar keywords
                    if has_similar_keywords(neighbor_doc.page_content, topic_keywords):
                        distance_boost *= 2.0
                    
                    # Special boost for department-specific content
                    if any(keyword in neighbor_doc.page_content.lower() 
                           for keyword in ["phòng thiết bị", "quản trị", "quân y"]):
                        distance_boost *= 2.5
                    
                    doc_chunks[neighbor_key]['boosted_score'] += distance_boost * position_boost
    
    # Sort by combined score (boosted_score + relevance_score), then by original ranking
    boosted_items = list(doc_chunks.values())
    boosted_items.sort(key=lambda x: (-(x['boosted_score'] + x['relevance_score']), x['original_index']))
    
    # Create final result list maintaining original documents not in doc_chunks
    boosted_docs = [item['doc'] for item in boosted_items]
    
    # Add any documents that weren't processed (no chunk_index)
    processed_docs = {id(item['doc']) for item in doc_chunks.values()}
    remaining_docs = [doc for doc in documents if id(doc) not in processed_docs]
    
    return boosted_docs + remaining_docs


def extract_topic_keywords(content: str, query: str) -> List[str]:
    """Extract topic-related keywords from content and query - generalized approach"""
    keywords = []
    
    # Extract meaningful phrases from query (remove stop words and short words)
    query_words = query.lower().split()
    meaningful_words = [word for word in query_words 
                       if len(word) > 2 and word not in ['của', 'là', 'có', 'được', 'trong', 'về', 'cho', 'từ', 'và', 'hoặc', 'các', 'một', 'này', 'đó']]
    keywords.extend(meaningful_words)
    
    # Extract section headers from content (numbered sections and Roman numerals)
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        # Look for numbered sections: 1., 2., 3., I., II., III., etc.
        if any(pattern in line for pattern in [
            '1.', '2.', '3.', '4.', '5.',
            'I.', 'II.', 'III.', 'IV.', 'V.',
            'VI.', 'VII.', 'VIII.', 'IX.', 'X.',
            'XI.', 'XII.', 'XIII.', 'XIV.', 'XV.'
        ]) and len(line) < 200:  # Avoid very long lines
            keywords.append(line.lower())
    
    # Extract department/organization names from content
    dept_patterns = ['phòng ', 'ban ', 'viện ', 'trung tâm ', 'khoa ']
    for line in lines:
        line_lower = line.lower()
        for pattern in dept_patterns:
            if pattern in line_lower and len(line) < 100:
                # Extract the department name
                start_idx = line_lower.find(pattern)
                dept_part = line[start_idx:start_idx+50].strip()
                if dept_part:
                    keywords.append(dept_part.lower())
    
    return list(set(keywords))  # Remove duplicates


def has_similar_keywords(content: str, keywords: List[str]) -> bool:
    """Check if content has similar keywords"""
    if not keywords:
        return False
    
    content_lower = content.lower()
    matches = sum(1 for keyword in keywords if keyword in content_lower)
    return matches >= max(1, len(keywords) * 0.3)  # At least 30% keyword overlap


def create_vector_database(output_path, data_dir="./data"):
    try:
        regulations = read_all_text_files(data_dir)

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400, 
            chunk_overlap=200, 
            length_function=len,
            separators=["\n\n", "\n", " ", ""],  
            keep_separator=False
        )
        chunks = text_splitter.split_text(regulations)

        embeddings = OllamaEmbeddings(
            model="nomic-embed-text",
            base_url=OLLAMA_BASE_URL
        )
        # embeddings = OllamaEmbeddings(
        #     model="nomic-embed-text"
        # )

        vectorstore = FAISS.from_texts(chunks, embeddings)

        if os.path.dirname(output_path):
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        else:
            os.makedirs(output_path, exist_ok=True)

        vectorstore.save_local(output_path)
        return chunks
    except Exception as e:
        print(f"Error creating vector database: {e}")
        raise


def load_vector_database(output_path, data_dir="./data"):
    try:
        embeddings = OllamaEmbeddings(
            model="nomic-embed-text",
            base_url=OLLAMA_BASE_URL
        )
        # embeddings = OllamaEmbeddings(
        #     model="nomic-embed-text"
        # )

        if not os.path.exists(output_path):
            chunks = create_vector_database(output_path, data_dir)
        else:
            regulations = read_all_text_files(data_dir)
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=400, 
                chunk_overlap=200, 
                length_function=len,
                separators=["\n\n", "\n", " ", ""], 
                keep_separator=False
            )
            chunks = text_splitter.split_text(regulations)

        print("Loading vector database...")
        return FAISS.load_local(output_path, embeddings, allow_dangerous_deserialization=True), chunks
    except Exception as e:
        print(f"Error loading vector database: {e}")
        raise


def create_hybrid_retriever(vector_db_path, data_dir="./data"):
    vectorstore, documents = load_vector_database(vector_db_path, data_dir)
    bm25_retriever = BM25Retriever.from_texts(texts=documents, k=15)

    return HybridRetriever(
        vectorstore=vectorstore, 
        bm25_retriever=bm25_retriever, 
        k=15
    ), documents


# Enhanced functions - preferred for new implementations
def create_vector_database_enhanced(output_path, data_dir="./data"):
    """Enhanced wrapper that returns enhanced documents"""
    documents = create_enhanced_vector_database(output_path, data_dir)
    return [doc.page_content for doc in documents]


def load_vector_database_enhanced(output_path, data_dir="./data"):
    """Enhanced wrapper that returns enhanced vectorstore and documents"""
    vectorstore, documents = load_enhanced_vector_database(output_path, data_dir)
    chunks = [doc.page_content for doc in documents]
    return vectorstore, chunks


def create_hybrid_retriever_enhanced(vector_db_path, data_dir="./data", window_size=1):
    """Enhanced wrapper that returns enhanced hybrid retriever with sliding window"""
    return create_enhanced_hybrid_retriever(vector_db_path, data_dir, window_size)


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    
    # Remove excessive whitespace and normalize line breaks
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Remove leading/trailing whitespace
        line = line.strip()
        
        # Skip empty lines but preserve paragraph breaks
        if not line:
            # Only add empty line if previous line wasn't empty
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
            
        # Remove excessive spaces within line
        line = re.sub(r'\s+', ' ', line)
        
        # Remove special characters that might interfere with processing
        # Keep Vietnamese characters, numbers, punctuation
        line = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\"\'\/\\\+\=\*\&\%\$\#\@\n\u00C0-\u1EF9]', '', line)
        
        cleaned_lines.append(line)
    
    # Join lines and normalize paragraph breaks
    cleaned_text = '\n'.join(cleaned_lines)
    
    # Remove multiple consecutive empty lines
    cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)
    
    return cleaned_text.strip()


def detect_and_preserve_structured_content(content: str, chunk_settings: Dict[str, Any]) -> List[str]:
    """Detect and preserve structured content like numbered lists, sections"""
    chunks = []
    
    # Pattern for structured lists (a), b), c), d), đ), e) or 1., 2., 3., etc.
    structured_patterns = [
        r'\b[a-zA-ZđĐ]\)\s+[^\n]{10,}',  # a), b), c), đ) patterns
        r'\b\d+\.\s+[^\n]{10,}',         # 1., 2., 3. patterns  
        r'\b[IVX]+\.\s+[^\n]{10,}',      # I., II., III. patterns
        r'-\s+[^\n]{10,}',               # - bullet points
    ]
    
    # Find all structured sections
    structured_sections = find_structured_sections(content, structured_patterns)
    
    if not structured_sections:
        return []  # No structured content found
    
    # Process content with structured sections preserved
    current_pos = 0
    
    for section_start, section_end, section_content in structured_sections:
        # Add content before structured section
        if section_start > current_pos:
            before_content = content[current_pos:section_start].strip()
            if before_content and len(before_content) > 50:
                before_chunks = enhanced_text_chunking(before_content, chunk_settings)
                chunks.extend(before_chunks)
        
        # Add structured section as complete chunk with context
        context_before = content[max(0, section_start-200):section_start].strip()
        context_after = content[section_end:min(len(content), section_end+200)].strip()
        
        complete_section = ""
        if context_before and not any(pattern in context_before for pattern in ['a)', 'b)', 'c)', 'd)', 'đ)', 'e)']):
            complete_section += context_before + "\n\n"
        
        complete_section += section_content
        
        if context_after and not any(pattern in context_after for pattern in ['a)', 'b)', 'c)', 'd)', 'đ)', 'e)']):
            complete_section += "\n\n" + context_after
        
        # Only split if the section is too large (over 2000 chars)
        if len(complete_section) > 2000:
            # Split while trying to preserve structure
            section_chunks = split_large_structured_section(complete_section, chunk_settings)
            chunks.extend(section_chunks)
        else:
            chunks.append(complete_section)
        
        current_pos = section_end
    
    # Add remaining content
    if current_pos < len(content):
        remaining_content = content[current_pos:].strip()
        if remaining_content:
            remaining_chunks = enhanced_text_chunking(remaining_content, chunk_settings)
            chunks.extend(remaining_chunks)
    
    return chunks if chunks else []


def find_structured_sections(content: str, patterns: List[str]) -> List[tuple]:
    """Find sections with structured content like lists"""
    import re
    
    sections = []
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Check if line matches any structured pattern
        for pattern in patterns:
            if re.match(pattern, line):
                # Found start of structured section
                section_start_line = i
                section_lines = [line]
                
                # Look for continuation of the same pattern type
                j = i + 1
                consecutive_items = 1
                
                while j < len(lines):
                    next_line = lines[j].strip()
                    
                    # Skip empty lines
                    if not next_line:
                        j += 1
                        continue
                    
                    # Check if continues the pattern
                    if re.match(pattern, next_line):
                        section_lines.append(next_line)
                        consecutive_items += 1
                        j += 1
                    else:
                        # Check if it's continuation of previous item (no pattern but indented or related)
                        if (len(next_line) > 20 and consecutive_items >= 2 and 
                            (next_line.startswith(' ') or next_line.startswith('\t') or 
                             any(keyword in next_line.lower() for keyword in ['theo', 'của', 'trong', 'được', 'phải']))):
                            section_lines.append(next_line)
                            j += 1
                        else:
                            break
                
                # Only consider as structured if we have multiple items (2+)
                if consecutive_items >= 2:
                    section_content = '\n'.join(section_lines)
                    section_start_pos = content.find(section_lines[0])
                    section_end_pos = section_start_pos + len(section_content)
                    
                    sections.append((section_start_pos, section_end_pos, section_content))
                    i = j  # Continue from after this section
                    break
                else:
                    i += 1
                    break
        else:
            i += 1
    
    # Sort sections by start position and merge overlapping
    sections.sort(key=lambda x: x[0])
    merged_sections = []
    
    for start, end, content in sections:
        if merged_sections and start <= merged_sections[-1][1] + 50:  # Allow small gap
            # Merge with previous section
            prev_start, prev_end, prev_content = merged_sections[-1]
            merged_content = prev_content + "\n" + content
            merged_sections[-1] = (prev_start, max(end, prev_end), merged_content)
        else:
            merged_sections.append((start, end, content))
    
    return merged_sections


def split_large_structured_section(section_content: str, chunk_settings: Dict[str, Any]) -> List[str]:
    """Split large structured sections while preserving related items"""
    lines = section_content.split('\n')
    chunks = []
    current_chunk = []
    current_size = 0
    max_size = chunk_settings.get('chunk_size', 1200)
    overlap_size = chunk_settings.get('chunk_overlap', 300)
    
    for line in lines:
        line_size = len(line)
        
        # If adding this line would exceed max size and we have content
        if current_size + line_size > max_size and current_chunk:
            # Complete current chunk
            chunks.append('\n'.join(current_chunk))
            
            # Start new chunk with overlap
            overlap_lines = []
            overlap_size_current = 0
            
            # Take last few lines for overlap
            for prev_line in reversed(current_chunk):
                if overlap_size_current + len(prev_line) <= overlap_size:
                    overlap_lines.insert(0, prev_line)
                    overlap_size_current += len(prev_line)
                else:
                    break
            
            current_chunk = overlap_lines
            current_size = overlap_size_current
        
        current_chunk.append(line)
        current_size += line_size
    
    # Add final chunk if any content remains
    if current_chunk:
        chunks.append('\n'.join(current_chunk))
    
    return chunks


def handle_table_content(content: str, chunk_settings: Dict[str, Any]) -> List[str]:
    """Handle table-containing content with enhanced processing"""
    chunks = []
    current_pos = 0
    
    while current_pos < len(content):
        # Look for table start
        table_start = content.find("[BẢNG DỮ LIỆU]", current_pos)
        
        if table_start == -1:
            # No more tables, chunk the rest normally
            remaining_content = content[current_pos:]
            if remaining_content.strip():
                remaining_chunks = enhanced_text_chunking(remaining_content, chunk_settings)
                chunks.extend(remaining_chunks)
            break
        
        # Add content before table if significant
        before_table = content[current_pos:table_start].strip()
        if len(before_table) > 100:
            before_chunks = enhanced_text_chunking(before_table, chunk_settings)
            chunks.extend(before_chunks)
        
        # Find table end and extract as complete unit
        table_end = content.find("[KẾT THÚC BẢNG]", table_start)
        if table_end == -1:
            remaining_content = content[table_start:]
            remaining_chunks = enhanced_text_chunking(remaining_content, chunk_settings)
            chunks.extend(remaining_chunks)
            break
        
        table_end += len("[KẾT THÚC BẢNG]")
        table_content = content[table_start:table_end]
        
        # Add context around table
        context_before = content[max(0, table_start-300):table_start].strip()
        context_after = content[table_end:min(len(content), table_end+300)].strip()
        
        complete_table = ""
        if context_before:
            complete_table += context_before + "\n\n"
        complete_table += table_content
        if context_after and not context_after.startswith("[BẢNG DỮ LIỆU]"):
            complete_table += "\n\n" + context_after
            current_pos = table_end + min(300, len(context_after))
        else:
            current_pos = table_end
        
        chunks.append(complete_table)
    
    return chunks


def enhanced_text_chunking(content: str, chunk_settings: Dict[str, Any]) -> List[str]:
    """Enhanced text chunking with improved overlap and separators"""
    # Enhanced separators that respect Vietnamese text structure
    separators = [
        "\n\n\n",     # Multiple paragraph breaks
        "\n\n",       # Paragraph breaks
        "\n",         # Line breaks
        ". ",         # Sentence endings with space
        ".",          # Sentence endings
        "; ",         # Semi-colons
        ", ",         # Commas with space  
        " ",          # Word boundaries
        "",           # Character level
    ]
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_settings.get('chunk_size', 1200),
        chunk_overlap=max(300, chunk_settings.get('chunk_overlap', 300)),  # Ensure minimum overlap
        length_function=len,
        separators=separators,
        keep_separator=True,  # Keep separators for better context
        is_separator_regex=False
    )
    
    chunks = text_splitter.split_text(content)
    
    # Post-process to ensure actual overlap exists
    if len(chunks) > 1:
        chunks = ensure_chunk_overlap(chunks, chunk_settings.get('chunk_overlap', 300))
    
    return chunks


def ensure_chunk_overlap(chunks: List[str], target_overlap: int) -> List[str]:
    """Ensure chunks have actual overlap by post-processing"""
    if len(chunks) <= 1:
        return chunks
    
    enhanced_chunks = []
    
    for i, chunk in enumerate(chunks):
        if i == 0:
            # First chunk remains as is
            enhanced_chunks.append(chunk)
        else:
            # Add overlap from previous chunk
            prev_chunk = chunks[i-1]
            
            # Take last 'target_overlap' characters from previous chunk
            overlap_text = prev_chunk[-target_overlap:] if len(prev_chunk) > target_overlap else prev_chunk
            
            # Find a good break point in the overlap (prefer word boundaries)
            if len(overlap_text) == target_overlap:
                # Find last space in the overlap to avoid splitting words
                last_space = overlap_text.rfind(' ')
                if last_space > target_overlap // 2:  # Only if we don't lose too much
                    overlap_text = overlap_text[last_space+1:]
            
            # Combine overlap with current chunk
            enhanced_chunk = overlap_text + "\n" + chunk
            enhanced_chunks.append(enhanced_chunk)
    
    return enhanced_chunks


def extract_table_data(doc) -> str:
    """Extract table data from Word document and format as text"""
    if not DOCX_AVAILABLE:
        return ""
    
    table_text = ""
    
    try:
        for table in doc.tables:
            table_text += "\n[BẢNG DỮ LIỆU]\n"
            
            # Extract table headers and data
            rows_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # Clean cell text
                    cell_text = re.sub(r'\s+', ' ', cell_text)
                    row_data.append(cell_text)
                rows_data.append(row_data)
            
            # Format table as structured text
            if rows_data:
                # Determine column widths for formatting
                max_cols = max(len(row) for row in rows_data) if rows_data else 0
                col_widths = [0] * max_cols
                
                for row in rows_data:
                    for i, cell in enumerate(row):
                        if i < max_cols:
                            col_widths[i] = max(col_widths[i], len(cell))
                
                # Format each row
                for i, row in enumerate(rows_data):
                    formatted_row = []
                    for j, cell in enumerate(row):
                        if j < max_cols:
                            # Pad cell to column width
                            padded_cell = cell.ljust(col_widths[j])
                            formatted_row.append(padded_cell)
                    
                    row_text = " | ".join(formatted_row)
                    table_text += row_text + "\n"
                    
                    # Add separator after header row
                    if i == 0 and len(rows_data) > 1:
                        separator = " | ".join(["-" * width for width in col_widths])
                        table_text += separator + "\n"
            
            table_text += "[KẾT THÚC BẢNG]\n\n"
    
    except Exception as e:
        print(f"Error extracting table data: {e}")
    
    return table_text


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from uploaded file - raises exceptions instead of returning error strings"""
    
    logger.info(f"🔍 Attempting to extract text from {file_path} (type: {file_type})")
    
    # Try docling first - handles most formats well
    if is_docling_available():
        try:
            logger.info("📄 Using Docling for extraction...")
            extracted_text = extract_text_with_docling(file_path)
            
            if extracted_text and extracted_text.strip():
                cleaned_text = clean_extracted_text(extracted_text)
                if cleaned_text:
                    logger.info(f"✅ Docling extraction successful: {len(cleaned_text)} characters")
                    return cleaned_text
                else:
                    logger.warning("⚠️ Docling returned content but it was empty after cleaning")
            
        except Exception as e:
            logger.warning(f"⚠️ Docling extraction failed: {str(e)}, trying fallback...")
    
    # Fallback methods for specific formats
    extracted_text = ""
    
    if file_type == "text/plain" or file_path.endswith('.txt'):
        logger.info("📝 Extracting .txt file...")
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            extracted_text = f.read()
    
    elif file_type == "application/pdf" or file_path.endswith('.pdf'):
        logger.info("📕 Extracting .pdf file...")
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing requires PyPDF2. Install with: pip install PyPDF2")
        
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
                    text_parts.append("\n")
        
        extracted_text = "".join(text_parts)
    
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_path.endswith('.docx'):
        logger.info("📘 Extracting .docx file...")
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing requires python-docx. Install with: pip install python-docx")
        
        doc = docx.Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                text_parts.append("\n")
        
        table_text = extract_table_data(doc)
        if table_text:
            text_parts.append(table_text)
        
        extracted_text = "".join(text_parts)
    
    elif file_type == "application/msword" or file_path.endswith('.doc'):
        logger.info("📗 Extracting legacy .doc file...")
        # Try python-docx first (some .doc files can be read as docx format)
        try:
            if DOCX_AVAILABLE:
                doc = docx.Document(file_path)
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                        text_parts.append("\n")
                extracted_text = "".join(text_parts)
            else:
                raise ValueError("No handler for .doc files available")
        except Exception as doc_error:
            error_msg = str(doc_error)
            logger.warning(f"python-docx failed for .doc: {error_msg}")
            
            # Check if this is actually an XML file (e.g., from office archives)
            if "is not a Word file" in error_msg or "xml" in error_msg.lower() or "openxml" in error_msg.lower():
                logger.error(f"❌ File appears to be XML or other format, not a valid .doc file: {error_msg}")
                raise ValueError(
                    f"❌ File format error: This file is not a valid .doc file.\n"
                    f"Content type: {file_type}\n"
                    f"Error: {error_msg}\n\n"
                    f"💡 Supported formats: .pdf, .docx, .doc, .txt, .md\n"
                    f"📝 Note: .doc files should be binary Word documents. "
                    f"If you have a .zip or .xml file, please upload the correct file format."
                )
            
            # For other errors, provide generic message
            raise ValueError(f"Cannot extract .doc file: {error_msg}")
    
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported formats: .txt, .pdf, .docx, .doc")
    
    # Clean and validate extracted text
    if not extracted_text:
        raise ValueError("File extraction produced no content - file may be empty or corrupted")
    
    cleaned_text = clean_extracted_text(extracted_text)
    
    if not cleaned_text or len(cleaned_text.strip()) == 0:
        raise ValueError("After processing, no readable text content found in the file")
    
    logger.info(f"✅ Text extraction successful: {len(cleaned_text)} characters extracted")
    return cleaned_text


def create_in_memory_retriever(file_content: str, chunk_size: int = 400, chunk_overlap: int = 200, k: int = 15) -> HybridRetriever:
    """Create an in-memory hybrid retriever from file content"""
    try:
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""],
            keep_separator=False
        )
        chunks = text_splitter.split_text(file_content)
        
      
        # embeddings = OllamaEmbeddings(
        #     model="nomic-embed-text"
        # )
        embeddings = OllamaEmbeddings(
            model="nomic-embed-text",
            base_url=OLLAMA_BASE_URL
        )
        # Create in-memory FAISS vector store
        vectorstore = FAISS.from_texts(chunks, embeddings)
        
        # Create BM25 retriever
        bm25_retriever = BM25Retriever.from_texts(texts=chunks, k=k)
        
        # Create hybrid retriever
        hybrid_retriever = HybridRetriever(
            vectorstore=vectorstore,
            bm25_retriever=bm25_retriever,
            k=k
        )
        
        return hybrid_retriever, chunks
    
    except Exception as e:
        print(f"Error creating in-memory retriever: {e}")
        raise


def process_uploaded_file(uploaded_file, chunk_size: int = 400, chunk_overlap: int = 200, k: int = 15):
    """Process an uploaded file and create an in-memory retriever"""
    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        # Extract text from file
        file_content = extract_text_from_file(tmp_file_path, uploaded_file.type)
        
        # Clean up temporary file
        os.unlink(tmp_file_path)
        
        if file_content.startswith("Error") or file_content.startswith("Unsupported") or "not available" in file_content:
            return None, file_content
        
        # Create in-memory retriever
        retriever, chunks = create_in_memory_retriever(file_content, chunk_size, chunk_overlap, k)
        
        return retriever, f"Successfully processed file: {uploaded_file.name}. Created {len(chunks)} chunks."
    
    except Exception as e:
        return None, f"Error processing file: {str(e)}"
