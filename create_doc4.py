# -*- coding: utf-8 -*-
"""Phase 4: Chương 4 + Kết luận + Tài liệu tham khảo + Phụ lục"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def fmt_run(run, bold=False, italic=False, size=13):
    run.font.name='Times New Roman'; run.font.size=Pt(size)
    run.font.bold=bold; run.font.italic=italic

def add_normal(doc, text, first_indent=True, bold=False):
    p = doc.add_paragraph(); p.style = doc.styles['Normal']
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.first_line_indent = Cm(1.0) if first_indent else Cm(0)
    run = p.add_run(text); fmt_run(run, bold=bold); return p

def add_heading1(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 1']
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(18); run.font.bold=True
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 2']
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(14); run.font.bold=True
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 3']
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(13); run.font.bold=True
    return p

def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'• {text}'); fmt_run(run, bold=bold); return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(12); run.font.italic=True
    return p

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); run._r.append(br)

def set_cell_bg(cell, color='BDD7EE'):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color)
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear')
    tcPr.append(shd)

def make_table(doc, headers, rows, col_widths=None, caption=None):
    if caption: add_caption(doc, caption)
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hrow = tbl.rows[0]
    for i,h in enumerate(headers):
        set_cell_bg(hrow.cells[i], 'BDD7EE')
        p=hrow.cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(h); run.font.name='Times New Roman'; run.font.size=Pt(11); run.font.bold=True
    for ri,row_data in enumerate(rows):
        row=tbl.rows[ri+1]
        for ci,val in enumerate(row_data):
            p=row.cells[ci].paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci==0 else WD_ALIGN_PARAGRAPH.LEFT
            run=p.add_run(str(val)); run.font.name='Times New Roman'; run.font.size=Pt(11)
    if col_widths:
        for ci,w in enumerate(col_widths):
            for row in tbl.rows: row.cells[ci].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return tbl

def add_ref(doc, num, text):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f'[{num}] {text}')
    run.font.name='Times New Roman'; run.font.size=Pt(12)
    return p

# ── Load part3 ────────────────────────────────────────────────────────────────
doc = Document('document/TAI_LIEU_SAN_PHAM_CHUAN_FORMAT_part3.docx')

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 4
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'CHƯƠNG 4. KẾT QUẢ VÀ THỰC NGHIỆM')

# 4.1
add_heading2(doc, '4.1. Môi trường thực nghiệm')
add_heading3(doc, '4.1.1. Cấu hình phần cứng')
add_normal(doc, 'Hệ thống được triển khai và đánh giá trên máy tính cá nhân với cấu hình phần cứng như sau:')
add_bullet(doc, 'CPU: Intel(R) Core(TM) i7-14700F')
add_bullet(doc, 'GPU: NVIDIA GeForce RTX 5060 Ti – VRAM: 16 GB')
add_bullet(doc, 'RAM: 32 GB')
add_bullet(doc, 'Ổ cứng: 1 TB')

add_heading3(doc, '4.1.2. Thông số Graph Builder và Retrieval')
add_normal(doc, 'Hệ thống sử dụng các thông số Graph Builder và Retrieval sau:')
add_bullet(doc, 'threshold (ngưỡng cosine tạo cạnh ngữ nghĩa) = 0.7')
add_bullet(doc, 'max_edges_per_node (số cạnh tối đa mỗi nút) = 5')
add_bullet(doc, 'embedding_model = "nomic-embed-text:latest" (triển khai qua Ollama)')
add_bullet(doc, 'Top-k documents = 10 (lấy ra 10 chunk liên quan nhất)')
add_bullet(doc, 'Routing communities = 3 (lấy ra 3 phòng ban liên quan nhất)')
add_bullet(doc, 'Metadata filter: Disabled')

add_heading3(doc, '4.1.3. Dataset thử nghiệm')
add_normal(doc, 'Để thử nghiệm hệ thống, bộ dữ liệu hành chính tại Học viện Kỹ thuật Mật mã được sử dụng. Các file dữ liệu được ban hành từ các phòng ban chính thức của Học viện.')

make_table(doc,
    ['Thông số', 'Giá trị', 'Mô tả'],
    [
        ['Số lượng document', '23', 'Tổng số file của các phòng ban'],
        ['Số Communities', '7', 'Số lượng phòng ban (cộng đồng đồ thị)'],
        ['Định dạng', 'PDF, DOCX, MD', 'Các định dạng tài liệu đầu vào'],
        ['Ngôn ngữ', 'Tiếng Việt', 'Ngôn ngữ chính của tài liệu'],
        ['Tổng số chunk', '[CẦN BỔ SUNG]', 'Tổng số đoạn văn bản sau chunking'],
        ['Kích thước chunk trung bình', '[CẦN BỔ SUNG]', 'Số token trung bình mỗi chunk'],
    ],
    col_widths=[5.0, 3.5, 6.0],
    caption='Bảng 4.1. Chi tiết bộ dữ liệu hành chính thử nghiệm'
)

add_normal(doc, 'Để đánh giá hệ thống, bộ dữ liệu đánh giá dạng CSV gồm 114 mẫu câu hỏi–câu trả lời tham chiếu được xây dựng thủ công, tập trung khai thác các tri thức đặc thù về quy chế, công tác khảo thí và các nghiệp vụ hành chính của Học viện.')

make_table(doc,
    ['Thông số', 'Giá trị', 'Mô tả'],
    [
        ['Số lượng mẫu', '114', 'Tổng số cặp câu hỏi–câu trả lời tham chiếu'],
        ['Loại câu hỏi', 'Factoid, Multi-hop', 'Câu hỏi đơn giản và câu hỏi liên đới nhiều điều khoản'],
        ['Nguồn tạo dữ liệu', 'Thủ công', 'Chuyên gia lĩnh vực hành chính tạo và kiểm duyệt'],
        ['Ngôn ngữ', 'Tiếng Việt', 'Câu hỏi và câu trả lời đều bằng tiếng Việt'],
    ],
    col_widths=[5.0, 3.5, 6.0],
    caption='Bảng 4.2. Mô tả bộ dữ liệu đánh giá'
)

# 4.2
add_heading2(doc, '4.2. Đánh giá mô hình Retrieval')
add_heading3(doc, '4.2.1. Các độ đo đánh giá')
add_normal(doc, 'Trong quá trình thực nghiệm, hệ thống được đánh giá bằng năm độ đo tiêu chuẩn trong lĩnh vực Information Retrieval và NLP:')
add_bullet(doc, 'Cosine Similarity: Đo độ tương đồng góc giữa hai vector. Giá trị càng gần 1 thì hai vector càng tương đồng về ngữ nghĩa.')
add_bullet(doc, 'Precision: Tỉ lệ kết quả liên quan trong tổng số kết quả trả về – đánh giá độ chính xác.')
add_bullet(doc, 'MAP (Mean Average Precision): Đánh giá chất lượng xếp hạng tổng hợp trên toàn bộ tập truy vấn.')
add_bullet(doc, 'MRR (Mean Reciprocal Rank): Đánh giá hệ thống có đưa kết quả đúng lên vị trí cao sớm hay không.')
add_bullet(doc, 'NDCG (Normalized Discounted Cumulative Gain): Độ đo xếp hạng mạnh nhất, xét đến mức độ liên quan của từng kết quả tại vị trí cụ thể.')

add_heading3(doc, '4.2.2. So sánh Vector RAG vs Graph RAG')
add_normal(doc, 'Bốn mô hình ngôn ngữ lớn được sử dụng trong thực nghiệm: Gemini 2.0 Flash API, Qwen3:32b, Qwen2.5:72b và Qwen2.5:7b. Kết quả được so sánh trên hai kiến trúc: RAG truyền thống và GraphRAG.')

make_table(doc,
    ['RAG + LLMs', 'Precision', 'Cosine Similarity', 'MAP', 'NDCG', 'MRR'],
    [
        ['Gemini 2.0 Flash', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
        ['Qwen3:32b', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
        ['Qwen2.5:72b', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
        ['Qwen2.5:7b', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
    ],
    col_widths=[4.0, 2.5, 3.5, 2.5, 2.5, 2.5],
    caption='Bảng 4.3. Kết quả thực nghiệm kiến trúc RAG kết hợp LLMs'
)

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Điền số liệu thực nghiệm đầy đủ từ Bảng 4.3 trong tài liệu gốc vào bảng trên]', first_indent=False)

make_table(doc,
    ['GraphRAG + LLMs', 'Precision', 'Cosine Similarity', 'MAP', 'NDCG', 'MRR'],
    [
        ['Gemini 2.0 Flash', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
        ['Qwen3:32b', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
        ['Qwen2.5:72b', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
        ['Qwen2.5:7b', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]', '[Xem tài liệu gốc]'],
    ],
    col_widths=[4.0, 2.5, 3.5, 2.5, 2.5, 2.5],
    caption='Bảng 4.4. Kết quả thực nghiệm kiến trúc GraphRAG kết hợp LLMs'
)

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Điền số liệu thực nghiệm đầy đủ từ Bảng 4.4 trong tài liệu gốc vào bảng trên]', first_indent=False)

add_normal(doc, 'Nhận xét kết quả: Mô hình Qwen2.5:72b là phù hợp nhất khi kết hợp với kiến trúc RAG truyền thống. Khi sử dụng kiến trúc GraphRAG, kết quả đều tốt hơn so với RAG thuần túy – thể hiện rõ nhất ở chỉ số MAP và NDCG. Các chỉ số Precision, MAP, NDCG khi sử dụng Gemini còn tốt hơn Qwen2.5:72b trong GraphRAG. Với các mô hình tham số nhỏ hơn (Qwen2.5:7b), GraphRAG giúp bù đắp đáng kể hạn chế của mô hình nhỏ trong xử lý câu hỏi phức tạp.')

# 4.3
add_heading2(doc, '4.3. Đánh giá hiệu năng hệ thống')
add_heading3(doc, '4.3.1. Response Time (Thời gian phản hồi)')
add_normal(doc, 'Hệ thống được đánh giá qua ba chỉ số thời gian phản hồi chính:')
add_bullet(doc, 'Độ trễ trung bình API thông thường: < 100 ms cho các yêu cầu không liên quan đến LLM (truy vấn dữ liệu tĩnh, thao tác quản lý). Kết quả này cho thấy backend và CSDL được tối ưu tốt.')
add_bullet(doc, 'Thời gian phản hồi mô hình ngôn ngữ: Hệ thống sử dụng streaming – thời gian đến token đầu tiên (TTFT) được duy trì ở mức chấp nhận được, cải thiện đáng kể trải nghiệm người dùng so với chờ toàn bộ câu trả lời.')
add_bullet(doc, 'Thời gian truy xuất RAG: Quá trình tìm kiếm trong FAISS/Qdrant hoàn thành trong khoảng 50–200 ms tùy kích thước corpus và số cộng đồng cần duyệt.')

add_heading3(doc, '4.3.2. Throughput và Scalability')
add_normal(doc, 'Sử dụng Apache JMeter và Locust để kiểm thử tải, hệ thống cho thấy khả năng xử lý đồng thời nhiều yêu cầu nhờ mô hình bất đồng bộ FastAPI kết hợp caching. Kết quả cho thấy hệ thống đáp ứng tốt trong các kịch bản sử dụng thực tế với lưu lượng truy cập cao.')
add_normal(doc, 'Về khả năng mở rộng, hệ thống được triển khai container hóa qua Docker Compose, cho phép dễ dàng nhân bản (scale out) Backend và dịch vụ Chatbot độc lập. Kiến trúc tách biệt Frontend – Backend – Chatbot AI đảm bảo mỗi thành phần có thể mở rộng độc lập theo nhu cầu.')

# 4.4
add_heading2(doc, '4.4. Demo chức năng hệ thống')
add_heading3(doc, '4.4.1. Chức năng Chat với AI')
add_normal(doc, 'Người dùng nhập câu hỏi vào khung chat. Hệ thống tiếp nhận truy vấn, phân tích ngữ nghĩa, định tuyến luồng xử lý phù hợp, truy xuất tài liệu qua Graph RAG và sinh câu trả lời dưới dạng streaming.')
add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 4.1 – Giao diện khung chat và thao tác gửi câu hỏi]', first_indent=False)
add_normal(doc, 'Chatbot sinh câu trả lời bám sát nội dung tài liệu nội bộ, hiển thị từng ký tự theo dạng streaming tạo trải nghiệm tự nhiên như đang gõ phím. Câu trả lời kèm trích dẫn nguồn tài liệu tương ứng.')
add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 4.2 – Kết quả phản hồi của Chatbot bám sát tài liệu nội bộ]', first_indent=False)

add_heading3(doc, '4.4.2. Quản lý lịch sử hội thoại')
add_normal(doc, 'Mỗi phiên trò chuyện được tự động lưu và hiển thị tại thanh điều hướng bên trái. Người dùng có thể truy cập lại, đổi tên từng phiên để phân loại, hoặc xóa các phiên không cần thiết.')
add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 4.3 – Giao diện danh sách lịch sử hội thoại]', first_indent=False)

add_heading3(doc, '4.4.3. Quản lý kho tri thức RAG (Admin)')
add_normal(doc, 'Admin tải lên tài liệu đầu vào (PDF, Word, TXT) theo thư mục phòng ban. Hệ thống tự động thực hiện: đọc và trích xuất nội dung → chia nhỏ văn bản thành chunks → chuyển đổi sang vector ngữ nghĩa → lưu trữ vào CSDL → rebuild đồ thị tri thức. Admin theo dõi danh sách tài liệu và thực hiện cập nhật khi có thay đổi.')
add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 4.4 – Giao diện tải file và quản lý tài liệu tri thức]', first_indent=False)

add_heading3(doc, '4.4.4. Thống kê và Quản trị hệ thống (Admin)')
add_normal(doc, 'Dashboard Admin cung cấp các chỉ số quan trọng theo thời gian thực: tổng số người dùng, số tin nhắn theo ngày/tháng, trạng thái API và mô hình AI. Admin có thể linh hoạt chuyển đổi giữa các mô hình ngôn ngữ và cấu hình giới hạn sử dụng.')
add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 4.5 – Dashboard thống kê hoạt động của Chatbot]', first_indent=False)

# 4.5
add_heading2(doc, '4.5. Phân tích và thảo luận')
add_heading3(doc, '4.5.1. Ưu điểm của hệ thống')
add_bullet(doc, 'Chất lượng suy luận và phản hồi được cải thiện rõ rệt: Kiến trúc RAG lai kết hợp tìm kiếm vector ngữ nghĩa và khai thác đồ thị tri thức giúp câu trả lời bám sát tài liệu, giảm hallucination xuống còn 1–2%, vượt trội so với LLM thuần túy (3–8%).')
add_bullet(doc, 'Kiến trúc hiện đại và có khả năng mở rộng: FastAPI bất đồng bộ kết hợp ReactJS đạt hiệu suất tốt, container hóa Docker cho phép triển khai và mở rộng linh hoạt.')
add_bullet(doc, 'Phù hợp với tiếng Việt: Hệ thống được thiết kế với dữ liệu và ngữ cảnh hướng tới môi trường tiếng Việt, sử dụng mô hình embedding hỗ trợ đa ngôn ngữ.')
add_bullet(doc, 'Đóng góp học thuật: 02 bài báo khoa học đã được công bố tại hội nghị quốc tế (ICSC 2025) và hội thảo quốc gia, xác nhận giá trị nghiên cứu của đề tài.')

add_heading3(doc, '4.5.2. Hạn chế và hướng cải thiện')
add_bullet(doc, 'Chi phí tài nguyên phần cứng còn cao: Vận hành Local LLM đòi hỏi GPU có VRAM lớn, hạn chế khả năng triển khai rộng rãi. Giải pháp: áp dụng lượng tử hóa INT8/INT4 để giảm yêu cầu phần cứng.')
add_bullet(doc, 'Chất lượng đồ thị phụ thuộc dữ liệu đầu vào: Tài liệu scan chất lượng thấp hoặc định dạng không chuẩn ảnh hưởng đến chất lượng chunking và đồ thị. Giải pháp: bổ sung bước tiền xử lý OCR nâng cao.')
add_bullet(doc, 'Khả năng xử lý bảng biểu phức tạp còn hạn chế: Một số bảng đặc thù trong văn bản pháp quy Việt Nam chưa được trích xuất hoàn toàn chính xác.')

# 4.6
add_heading2(doc, '4.6. Tổng kết chương 4')
add_normal(doc, 'Chương 4 đã trình bày đầy đủ quá trình thực nghiệm và đánh giá hệ thống KMA Agent trên ba phương diện: chất lượng truy xuất thông tin, hiệu năng vận hành và tính năng thực tế. Kết quả thực nghiệm khẳng định kiến trúc GraphRAG vượt trội so với RAG truyền thống trên bộ dữ liệu hành chính tiếng Việt, đặc biệt với các câu hỏi đòi hỏi suy luận liên đới nhiều điều khoản. Hệ thống đạt độ trễ phản hồi thấp, xử lý đồng thời ổn định và các chức năng đều được triển khai và kiểm thử thành công.')

# ═══════════════════════════════════════════════════════════════════════════════
# KẾT LUẬN
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'KẾT LUẬN')

add_heading2(doc, 'Kết quả đạt được')
add_normal(doc, 'Sau quá trình nghiên cứu và triển khai, đề tài "Xây dựng hệ thống Chatbot hỗ trợ truy vấn điểm của sinh viên sử dụng kỹ thuật RAG" đã hoàn thành với kết quả xuất sắc trên nhiều phương diện.')
add_normal(doc, 'Về mặt thực tiễn và khoa học, quá trình nghiên cứu đã đóng góp trực tiếp vào cộng đồng học thuật với 02 bài báo khoa học:')
add_bullet(doc, '"An Enhanced Chatbot Architecture via Generative AI and GraphRAG" – Hội nghị quốc tế lần thứ 8 (ICSC 2025), Chỉ số SJR Q2.')
add_bullet(doc, '"Developing a Virtual Assistant for Exam Score Retrieval Using a Local LLM" – Hội thảo quốc gia lần thứ XXVIII về một số vấn đề chọn lọc của CNTT và TT, tháng 4/2026.')
add_normal(doc, 'Về mặt hệ thống và giải pháp công nghệ:')
add_bullet(doc, 'Xây dựng thành công nền tảng Chatbot thông minh phục vụ tra cứu thông tin học vụ, quy chế đào tạo, thủ tục hành chính và tài liệu nội bộ tại Học viện Kỹ thuật Mật mã.')
add_bullet(doc, 'Áp dụng thành công kiến trúc Multi-Agent System (Supervisor Agent – Worker Agent) kết hợp Graph RAG và Hybrid Retrieval, vượt trội so với Vector RAG truyền thống.')
add_bullet(doc, 'Tích hợp cơ chế ReAct Agent xử lý hội thoại đa lượt linh hoạt; Backend FastAPI bất đồng bộ + ReactJS Frontend + Docker đảm bảo hiệu năng và khả năng mở rộng.')
add_bullet(doc, 'Hệ thống góp phần thúc đẩy chuyển đổi số tại Học viện, giảm tải công việc cho các phòng ban chức năng, đặc biệt chức năng tra cứu điểm theo thời gian thực.')

add_heading2(doc, 'Hạn chế và hướng phát triển')
add_normal(doc, 'Mặc dù đạt được nhiều kết quả tích cực, hệ thống còn tồn tại một số hạn chế cần tiếp tục nghiên cứu khắc phục:')
add_normal(doc, 'Hạn chế:', bold=True, first_indent=False)
add_bullet(doc, 'Chi phí tài nguyên phần cứng để vận hành Local LLM còn tương đối cao, đặc biệt ở quy mô dữ liệu lớn và số lượng người dùng đồng thời cao.')
add_bullet(doc, 'Chất lượng Knowledge Graph phụ thuộc vào bước tiền xử lý tài liệu (OCR, chunking văn bản pháp quy tiếng Việt và phân tích bảng biểu phức tạp).')
add_bullet(doc, 'Khả năng xử lý tài liệu scan chất lượng thấp và bảng biểu đặc thù chưa hoàn toàn chính xác.')
add_normal(doc, 'Hướng phát triển tiếp theo:', bold=True, first_indent=False)
add_bullet(doc, 'Tối ưu hóa mô hình: Áp dụng lượng tử hóa (Quantization – INT4/INT8), chưng cất tri thức (Knowledge Distillation) và triển khai mô hình lai (Local + Cloud API) để cân bằng chi phí – hiệu suất.')
add_bullet(doc, 'Nâng cấp dữ liệu và Graph RAG: Mở rộng, làm sạch và chuẩn hóa dataset nội bộ; nghiên cứu kỹ thuật Entity Linking và Relation Extraction nâng cao.')
add_bullet(doc, 'Mở rộng tính năng: Tích hợp voice interaction (chuyển văn bản thành giọng nói), phát triển Mobile App (React Native – Android/iOS), hỗ trợ đa phương thức.')
add_bullet(doc, 'Đánh giá toàn diện hơn: Xây dựng bộ dữ liệu đánh giá đặc thù cho môi trường giáo dục Việt Nam; thực hiện A/B testing với người dùng thực tế.')

# ═══════════════════════════════════════════════════════════════════════════════
# TÀI LIỆU THAM KHẢO
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'TÀI LIỆU THAM KHẢO')

refs = [
    'A. Vaswani et al., "Attention Is All You Need," in Advances in Neural Information Processing Systems (NeurIPS), vol. 30, 2017, pp. 5998–6008.',
    'M. Raza, Z. Jahangir, M. B. Riaz, M. J. Saeed, and M. A. Sattar, "Industrial Applications of Large Language Models," Scientific Reports, vol. 15, 2025.',
    'T. B. Brown et al., "Language Models Are Few-Shot Learners," in Advances in Neural Information Processing Systems, vol. 33, 2020, pp. 1877–1901.',
    'H. Touvron et al., "LLaMA: Open and Efficient Foundation Language Models," arXiv preprint arXiv:2302.13971, 2023.',
    'S. Roller et al., "Recipes for Building an Open-Domain Chatbot," in Proceedings of EACL 2021.',
    'M. Zhong, Y. Liu, Y. Xu, C. Zhu, and M. Zeng, "DialogLM: Pre-trained Model for Long Dialogue Understanding and Summarization," in AAAI, 2022.',
    'OpenAI, "ChatGPT: Optimizing Language Models for Dialogue," 2022. [Trực tuyến]. Available: https://openai.com/blog/chatgpt.',
    'Anthropic, "Introducing Claude," 2023. [Trực tuyến]. Available: https://www.anthropic.com/index/introducing-claude.',
    'A. Dubey et al., "The Llama 3 Herd of Models," arXiv preprint arXiv:2407.21783, 2024.',
    'L. Huang, J. Lin, and Y. Chen, "FastChat: Open Platform for Training, Serving, and Evaluating Large Language Models," 2023.',
    'W. Chen, X. Zhang, and Y. Li, "Applying GPT-4 in education: Automated grading and feedback generation," in Proc. Int. Conf. Comput. Educ., 2023.',
    'S. Bubeck et al., "Sparks of artificial general intelligence: Early experiments with GPT-4," arXiv preprint arXiv:2303.12712, 2023.',
    'OpenAI et al., "GPT-4 Technical Report," arXiv preprint arXiv:2303.08774, 2023.',
    'P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," in Advances in NeurIPS, vol. 33, 2020, pp. 9459–9474.',
    'W. Yu, "Retrieval-augmented generation across heterogeneous knowledge," in Proc. NAACL 2022.',
    'K. Pichai, "A retrieval-augmented generation based large language model benchmarked on a novel dataset," Journal of Student Research, vol. 12, no. 4, 2023.',
    'J. Gu, "A Research of Challenges and Solutions in Retrieval-Augmented Generation (RAG) Systems," Highlights in Science, Engineering and Technology, 2024.',
    'D. Edge et al., "From local to global: A Graph RAG approach to query-focused summarization," arXiv preprint arXiv:2404.16130, 2024.',
    'H. Han et al., "RAG vs. GraphRAG: A Systematic Evaluation and Key Insights," arXiv preprint, 2025.',
    'S. Knollmeyer, O. Caymazer, and D. Grossmann, "Document GraphRAG: Knowledge Graph Enhanced RAG for Document Question Answering," 2024.',
    'T. Zang et al., "Event-State Knowledge Graph-Enhanced GraphRAG for Order-Driven 3D Printing Supply Chain Configuration," Expert Systems with Applications, 2025.',
    'Z. Ji et al., "Survey of Hallucination in Natural Language Generation," ACM Computing Surveys, vol. 55, no. 12, 2023.',
    'V. Mavroudis, "LangChain," Preprints, 2024.',
    'T. Taulli and G. Deshmukh, "Introduction to LangGraph," in Building Generative AI Agents, Springer, 2025, pp. 209–235.',
    'S. Veturi et al., "RAG based Question-Answering for Contextual Response Prediction System," arXiv preprint, 2024.',
    'ductrong203, "kma_agent," GitHub. [Trực tuyến]. Available: https://github.com/ductrong203/kma_agent.git.',
    'V. H. Pham, D. T. Le, D. T. Bui, and T. H. V. Le, "An Enhanced Chatbot Architecture via Generative AI and GraphRAG," in Proc. 8th Int. Conf. Smart Computing and Communication (ICSC 2025), 2025.',
    'T. D. Le, T. D. Bui, D. T. Nguyen, and L. T. Bui, "Developing a Virtual Assistant for Exam Score Retrieval Using a Local LLM," in Proc. National Conf. 28th, Hanoi, April 2026.',
]

for i, ref in enumerate(refs, 1):
    add_ref(doc, i, ref)

# ═══════════════════════════════════════════════════════════════════════════════
# PHỤ LỤC
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'PHỤ LỤC')

add_heading2(doc, 'Phụ lục A. Danh sách từ viết tắt bổ sung')
add_normal(doc, 'Xem bảng Danh mục từ viết tắt ở phần đầu tài liệu.')

add_heading2(doc, 'Phụ lục B. Hướng dẫn cài đặt và triển khai')
add_normal(doc, 'Mã nguồn hệ thống được lưu trữ tại:')
add_bullet(doc, 'GitHub: https://github.com/ductrong203/kma_agent.git')
add_bullet(doc, 'Google Drive (tài liệu bổ sung): https://drive.google.com/drive/folders/1NDjhPuqgMJnqKfEhPDqLieRI7NVGByxd')

add_normal(doc, 'Yêu cầu môi trường:')
add_bullet(doc, 'Python 3.10+, Node.js 18+')
add_bullet(doc, 'Docker & Docker Compose')
add_bullet(doc, 'NVIDIA GPU với VRAM ≥ 16 GB (để chạy Local LLM qua Ollama)')
add_bullet(doc, 'RAM ≥ 16 GB, Ổ cứng ≥ 50 GB trống')

add_normal(doc, 'Các bước triển khai cơ bản:')
add_bullet(doc, 'Clone repository: git clone https://github.com/ductrong203/kma_agent.git')
add_bullet(doc, 'Cấu hình biến môi trường: sao chép .env.example thành .env và điền thông tin kết nối CSDL, API key')
add_bullet(doc, 'Khởi động dịch vụ: docker-compose up -d')
add_bullet(doc, 'Tải mô hình embedding: ollama pull nomic-embed-text:latest')
add_bullet(doc, 'Tải tài liệu RAG qua Admin Dashboard và thực hiện rebuild index')

add_heading2(doc, 'Phụ lục C. Danh sách bài báo khoa học liên quan')
add_normal(doc, 'Đề tài đã đóng góp vào cộng đồng khoa học với 02 bài báo được công bố:')
add_bullet(doc, '[27] V. H. Pham, D. T. Le, D. T. Bui, and T. H. V. Le, "An Enhanced Chatbot Architecture via Generative AI and GraphRAG," in Proc. 8th International Conference on Smart Computing and Communication (ICSC 2025), 2025. (Chỉ số SJR Q2)')
add_bullet(doc, '[28] T. D. Le, T. D. Bui, D. T. Nguyen, and L. T. Bui, "Developing a Virtual Assistant for Exam Score Retrieval Using a Local LLM," in Proc. National Conference 28th, Hanoi, April 2026.')

# ── Save final ─────────────────────────────────────────────────────────────────
doc.save('document/TAI_LIEU_SAN_PHAM_CHUAN_FORMAT.docx')
print("✅ TAI_LIEU_SAN_PHAM_CHUAN_FORMAT.docx saved successfully!")
