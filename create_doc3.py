# -*- coding: utf-8 -*-
"""Phase 3: Chương 3 – Phân tích, Thiết kế và Cài đặt"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def fmt_run(run, bold=False, italic=False, size=13):
    run.font.name='Times New Roman'; run.font.size=Pt(size)
    run.font.bold=bold; run.font.italic=italic

def add_normal(doc, text, first_indent=True, bold=False):
    p = doc.add_paragraph(); p.style = doc.styles['Normal']
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.first_line_indent = Cm(1.0) if first_indent else Cm(0)
    run = p.add_run(text); fmt_run(run, bold=bold); return p

def add_heading1(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 1']
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(18); run.font.bold=True
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 2']
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(14); run.font.bold=True
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 3']
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(13); run.font.bold=True
    return p

def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'• {text}'); fmt_run(run, bold=bold); return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(12); run.font.italic=True
    return p

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); run._r.append(br)

def set_cell_bg(cell, color='BDD7EE'):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear')
    tcPr.append(shd)

def make_table(doc, headers, rows, col_widths=None, caption=None):
    if caption: add_caption(doc, caption)
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hrow = tbl.rows[0]
    for i,h in enumerate(headers):
        set_cell_bg(hrow.cells[i], 'BDD7EE')
        p=hrow.cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run(h); run.font.name='Times New Roman'; run.font.size=Pt(11); run.font.bold=True
    for ri,row_data in enumerate(rows):
        row=tbl.rows[ri+1]
        for ci,val in enumerate(row_data):
            p=row.cells[ci].paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci==0 else WD_ALIGN_PARAGRAPH.LEFT
            run=p.add_run(str(val)); run.font.name='Times New Roman'; run.font.size=Pt(11)
    if col_widths:
        for ci,w in enumerate(col_widths):
            for row in tbl.rows: row.cells[ci].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return tbl

def usecase_table(doc, uc_id, uc_name, actor, desc, precond, result, main_flow, alt_flow, rule='', note=''):
    add_caption(doc, f'Bảng {uc_id}. Bảng đặc tả usecase {uc_name.lower()}')
    tbl = doc.add_table(rows=8, cols=2); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    fields = [
        ('Mã Use Case', uc_id.split('.')[1].strip() if '.' in uc_id else uc_id),
        ('Tên Use Case', uc_name),
        ('Actor', actor),
        ('Mô tả ngắn gọn', desc),
        ('Điều kiện tiên quyết', precond),
        ('Kết quả sau thực hiện', result),
        ('Luồng chính', main_flow),
        ('Luồng thay thế / ngoại lệ', alt_flow),
    ]
    for ri, (label, val) in enumerate(fields):
        row = tbl.rows[ri]
        set_cell_bg(row.cells[0], 'D9E1F2')
        p0=row.cells[0].paragraphs[0]; p0.alignment=WD_ALIGN_PARAGRAPH.LEFT
        r0=p0.add_run(label); r0.font.name='Times New Roman'; r0.font.size=Pt(11); r0.font.bold=True
        p1=row.cells[1].paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.LEFT
        r1=p1.add_run(val); r1.font.name='Times New Roman'; r1.font.size=Pt(11)
    for row in tbl.rows:
        row.cells[0].width=Cm(4.5); row.cells[1].width=Cm(11.0)
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

def db_table(doc, caption_text, headers, rows, col_widths):
    make_table(doc, headers, rows, col_widths, caption_text)

def api_table(doc, caption_text, endpoint, method, params, result):
    add_caption(doc, caption_text)
    data = [('Endpoint', endpoint), ('Phương thức', method),
            ('Tham số', params), ('Kết quả trả về', result)]
    tbl = doc.add_table(rows=len(data), cols=2); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,(label,val) in enumerate(data):
        row=tbl.rows[ri]
        set_cell_bg(row.cells[0],'D9E1F2')
        p0=row.cells[0].paragraphs[0]; r0=p0.add_run(label)
        r0.font.name='Times New Roman'; r0.font.size=Pt(11); r0.font.bold=True
        p1=row.cells[1].paragraphs[0]; r1=p1.add_run(val)
        r1.font.name='Times New Roman'; r1.font.size=Pt(11)
        row.cells[0].width=Cm(4.0); row.cells[1].width=Cm(11.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# ── Load part2 ────────────────────────────────────────────────────────────────
doc = Document('document/TAI_LIEU_SAN_PHAM_CHUAN_FORMAT_part2.docx')

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 3
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'CHƯƠNG 3. PHÂN TÍCH, THIẾT KẾ VÀ CÀI ĐẶT HỆ THỐNG')

# 3.1
add_heading2(doc, '3.1. Phân tích hệ thống')
add_heading3(doc, '3.1.1. Tổng quan và đối tượng sử dụng')
add_normal(doc, 'Hệ thống Chatbot AI đa mô hình được xây dựng với mục tiêu hỗ trợ tra cứu thông tin học vụ, quy chế đào tạo, thủ tục hành chính và tài liệu nội bộ tại Học viện Kỹ thuật Mật mã thông qua giao tiếp hội thoại tự nhiên. Hệ thống hướng tới hai nhóm người dùng chính:')
add_bullet(doc, 'Người dùng thông thường (sinh viên, giảng viên): Sử dụng Chatbot để tra cứu thông tin, đặt câu hỏi, quản lý lịch sử hội thoại và lưu trữ tài liệu cá nhân.')
add_bullet(doc, 'Quản trị viên (Admin): Quản lý tài khoản người dùng, cấu hình mô hình AI, quản lý dữ liệu RAG, thiết lập giới hạn hệ thống và theo dõi vận hành.')

add_heading3(doc, '3.1.2. Yêu cầu chức năng')
add_normal(doc, 'Hệ thống cần cung cấp các chức năng chính sau:')
add_bullet(doc, 'Quản lý tài khoản: Đăng ký, đăng nhập, xác thực phiên làm việc và phân quyền User/Admin.')
add_bullet(doc, 'Chat với AI: Tiếp nhận câu hỏi từ người dùng, truy xuất dữ liệu liên quan qua Graph RAG, gửi yêu cầu tới mô hình AI và trả về câu trả lời phù hợp.')
add_bullet(doc, 'Quản lý lịch sử hội thoại: Lưu lại các cuộc trò chuyện, cho phép xem lại, đổi tên hoặc xóa khi cần.')
add_bullet(doc, 'Quản lý file cá nhân: Tải lên, xem và xóa tài liệu cá nhân trong giới hạn cho phép.')
add_bullet(doc, 'Quản lý mô hình LLM (Admin): Thêm mới, chỉnh sửa cấu hình, bật/tắt trạng thái hoạt động của mô hình.')
add_bullet(doc, 'Quản lý dữ liệu RAG (Admin): Cập nhật tài liệu, xử lý embedding và xây dựng chỉ mục tìm kiếm.')
add_bullet(doc, 'Thiết lập giới hạn hệ thống (Admin): Cấu hình số lượt truy cập, dung lượng file, hạn mức token.')
add_bullet(doc, 'Thống kê sử dụng (Admin): Tổng hợp số lượng truy cập, mức sử dụng mô hình và các chỉ số vận hành.')

add_heading3(doc, '3.1.3. Yêu cầu phi chức năng')
add_normal(doc, 'Các yêu cầu phi chức năng đảm bảo hệ thống hoạt động ổn định và hiệu quả:')
add_bullet(doc, 'Hiệu năng: Thời gian phản hồi API thông thường < 100ms; chức năng chat AI trả kết quả theo từng phần (streaming) để cải thiện trải nghiệm.')
add_bullet(doc, 'Khả năng mở rộng: Hỗ trợ nhiều người dùng đồng thời, dễ mở rộng theo chiều ngang qua Docker Compose.')
add_bullet(doc, 'Bảo mật: Áp dụng JWT trong xác thực, phân quyền rõ ràng giữa User và Admin, mã hóa mật khẩu.')
add_bullet(doc, 'Độ ổn định: Hoạt động liên tục với cơ chế Rate Limiting và xử lý lỗi graceful.')
add_bullet(doc, 'Dễ bảo trì: Mã nguồn tổ chức theo module rõ ràng, hỗ trợ triển khai Docker.')

add_heading3(doc, '3.1.4. Biểu đồ use case tổng quát')
add_normal(doc, 'Hệ thống có hai tác nhân chính: Người dùng và Quản trị viên. Quản trị viên kế thừa toàn bộ quyền của người dùng và được bổ sung các chức năng quản trị. Ngoài ra, hệ thống tích hợp các thành phần ngoài: nhà cung cấp mô hình AI (Gemini API, Ollama), Vector Database (FAISS/Qdrant), và dịch vụ xác thực JWT.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 3.1 – Biểu đồ use case tổng quát]', first_indent=False)

add_heading3(doc, '3.1.5. Đặc tả use case')

usecase_table(doc, 'Bảng 3.1',
    'Đăng ký tài khoản', 'Người dùng',
    'Cho phép người dùng mới tạo tài khoản để truy cập hệ thống',
    'Người dùng chưa đăng nhập và đang ở trang đăng ký',
    'Tài khoản được tạo thành công, người dùng nhận thông báo thành công',
    '1. Người dùng mở form đăng ký\n2. Nhập email, mật khẩu, xác nhận mật khẩu\n3. Hệ thống kiểm tra tính hợp lệ\n4. Tạo tài khoản mới trong CSDL',
    'A1: Email đã tồn tại → Hiển thị thông báo "Email này đã được sử dụng"\nE1: Mật khẩu không khớp → Yêu cầu nhập lại')

usecase_table(doc, 'Bảng 3.2',
    'Đăng nhập', 'Người dùng, Admin',
    'Xác thực người dùng và điều hướng đến trang phù hợp theo vai trò',
    'Người dùng hoặc Admin đã có tài khoản',
    'JWT token được cấp; người dùng được chuyển đến trang chat hoặc trang quản trị',
    '1. Mở màn hình đăng nhập\n2. Nhập email và mật khẩu\n3. Hệ thống xác thực thông tin\n4. Tạo và trả về JWT access token + refresh token',
    'E1: Sai thông tin đăng nhập → Hiển thị thông báo lỗi\nE2: Tài khoản bị khóa → Thông báo liên hệ Admin',
    'Giới hạn số lần nhập sai để tránh tấn công brute force',
    'Hỗ trợ tùy chọn "Ghi nhớ đăng nhập"')

usecase_table(doc, 'Bảng 3.3',
    'Chat với AI', 'Người dùng',
    'Người dùng gửi câu hỏi và nhận câu trả lời từ mô hình AI thông qua cơ chế Graph RAG',
    'Người dùng đã đăng nhập vào hệ thống',
    'Câu trả lời được sinh và hiển thị theo dạng streaming',
    '1. Người dùng nhập câu hỏi\n2. Hệ thống phân tích ý định và định tuyến\n3. Truy xuất tài liệu liên quan qua Graph RAG\n4. Gửi ngữ cảnh đến LLM và nhận câu trả lời\n5. Hiển thị kết quả dạng streaming',
    'A1: Câu hỏi ngoài phạm vi → Thông báo giới hạn phạm vi\nE1: Vượt quá rate limit → Trả về thông báo giới hạn',
    'Áp dụng rate limit theo gói tài khoản; ưu tiên sử dụng RAG khi có file liên quan')

usecase_table(doc, 'Bảng 3.4',
    'Quản lý lịch sử hội thoại', 'Người dùng',
    'Xem, xóa hoặc chỉnh sửa tiêu đề các cuộc hội thoại trước đó',
    'Người dùng đã đăng nhập',
    'Danh sách cuộc trò chuyện được hiển thị hoặc cập nhật theo yêu cầu',
    '1. Người dùng mở danh sách hội thoại\n2. Chọn thao tác: xem / đổi tên / xóa\n3. Hệ thống thực hiện thao tác và cập nhật giao diện',
    'E1: Không có hội thoại nào → Hiển thị thông báo trống\nE2: Xóa thất bại → Thông báo lỗi')

usecase_table(doc, 'Bảng 3.5',
    'Quản lý file cá nhân', 'Người dùng',
    'Tải lên, xem, tải xuống hoặc xóa các tài liệu cá nhân',
    'Người dùng đã đăng nhập',
    'File được xử lý theo thao tác người dùng chọn',
    '1. Mở khu vực quản lý file\n2. Chọn thao tác (Tải lên / Xem / Xóa)\n3. Hệ thống xử lý file (embedding nếu tải lên mới)',
    'E1: File không hợp lệ (định dạng hoặc dung lượng vượt giới hạn) → Thông báo lỗi')

usecase_table(doc, 'Bảng 3.6',
    'Xem danh sách mô hình LLM', 'Người dùng',
    'Xem danh sách và thông tin chi tiết các mô hình AI đang khả dụng',
    'Người dùng đã đăng nhập',
    'Danh sách mô hình được hiển thị với trạng thái hoạt động',
    '1. Người dùng truy cập mục mô hình\n2. Hệ thống trả về danh sách mô hình đang hoạt động',
    'E1: Không có mô hình nào khả dụng → Thông báo liên hệ Admin')

usecase_table(doc, 'Bảng 3.7',
    'Quản lý mô hình LLM', 'Admin',
    'Thêm, sửa, xóa hoặc thay đổi trạng thái các mô hình LLM trong hệ thống',
    'Admin đã đăng nhập',
    'Mô hình được cập nhật và trạng thái phản ánh ngay lập tức',
    '1. Admin truy cập trang quản lý mô hình\n2. Chọn thao tác (Thêm / Sửa / Kích hoạt / Tắt)\n3. Hệ thống lưu cấu hình',
    'E1: Tên mô hình trùng → Thông báo lỗi',
    'Chỉ một mô hình được kích hoạt tại một thời điểm')

usecase_table(doc, 'Bảng 3.8',
    'Quản lý dữ liệu RAG', 'Admin',
    'Thêm, cập nhật hoặc xóa tài liệu dùng cho cơ chế Retrieval-Augmented Generation',
    'Admin đã đăng nhập',
    'Tài liệu được embedding và đồ thị tri thức được cập nhật',
    '1. Admin chọn thư mục phòng ban\n2. Tải lên hoặc xóa tài liệu\n3. Hệ thống tiền xử lý, chunking, tạo embedding\n4. Rebuild RAG index',
    'E1: File không đọc được → Thông báo lỗi\nE2: Rebuild thất bại → Ghi log và thông báo',
    'Tự động tạo embedding khi thêm hoặc cập nhật tài liệu')

usecase_table(doc, 'Bảng 3.9',
    'Quản lý người dùng', 'Admin',
    'Quản lý tài khoản người dùng: thêm, sửa, xóa, khóa/mở tài khoản',
    'Admin đã đăng nhập',
    'Tài khoản được cập nhật theo thao tác của Admin',
    '1. Admin truy cập trang quản lý người dùng\n2. Tìm kiếm người dùng\n3. Thực hiện thao tác (sửa vai trò / khóa / xóa)',
    'E1: Xóa tài khoản đang đăng nhập → Thông báo không thể thực hiện')

usecase_table(doc, 'Bảng 3.10',
    'Cấu hình giới hạn hệ thống', 'Admin',
    'Thiết lập các giới hạn sử dụng: rate limit, dung lượng file, số request',
    'Admin đã đăng nhập',
    'Cấu hình giới hạn mới được lưu và áp dụng ngay',
    '1. Truy cập trang cấu hình giới hạn\n2. Nhập các mức giới hạn (số request/phút, dung lượng tối đa)\n3. Hệ thống lưu và áp dụng',
    'E1: Giá trị nhập không hợp lệ → Hiển thị thông báo lỗi')

usecase_table(doc, 'Bảng 3.11',
    'Thống kê sử dụng', 'Admin',
    'Xem báo cáo thống kê sử dụng hệ thống theo thời gian',
    'Admin đã đăng nhập',
    'Dashboard thống kê được hiển thị với dữ liệu cập nhật',
    '1. Admin truy cập Dashboard\n2. Hệ thống tổng hợp: tổng người dùng, số tin nhắn, token sử dụng, trạng thái API',
    'E1: Dữ liệu thống kê trống → Hiển thị giá trị 0')

add_heading3(doc, '3.1.6. Biểu đồ hoạt động các chức năng chính')
for i, name in enumerate(['đăng ký','đăng nhập','chat với AI','quản lý lịch sử hội thoại',
                           'quản lý file cá nhân','xem danh sách mô hình','quản lý mô hình LLM',
                           'quản lý dữ liệu RAG','quản lý người dùng','giới hạn hệ thống','thống kê sử dụng'], 2):
    add_normal(doc, f'[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 3.{i} – Biểu đồ hoạt động chức năng {name}]', first_indent=False)

# 3.2
add_heading2(doc, '3.2. Thiết kế hệ thống')
add_heading3(doc, '3.2.1. Thiết kế kiến trúc tổng thể')
add_normal(doc, 'Kiến trúc tổng thể được thiết kế theo mô hình client-server tinh gọn với năm thành phần chính:')
add_bullet(doc, 'Frontend (ReactJS): Giao diện người dùng, tiếp nhận thao tác và hiển thị kết quả phản hồi.')
add_bullet(doc, 'Backend (FastAPI/Python): Trung tâm điều phối, tiếp nhận yêu cầu từ Frontend, truy xuất dữ liệu từ CSDL và giao tiếp với Chatbot AI.')
add_bullet(doc, 'CSDL (PostgreSQL + MongoDB + Qdrant): Lưu trữ dữ liệu điểm sinh viên (PostgreSQL), lịch sử hội thoại/tài khoản (MongoDB) và vector embedding (Qdrant).')
add_bullet(doc, 'Chatbot AI (Python + LangChain): Khối xử lý hội thoại thông minh, thực thi Graph RAG và sinh câu trả lời qua LLM.')
add_bullet(doc, 'Mobile App (React Native – hướng phát triển tương lai): Giao tiếp với Backend thông qua API đã xây dựng.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 3.13 – Kiến trúc tổng thể hệ thống]', first_indent=False)

add_heading3(doc, '3.2.2. Thiết kế kiến trúc Graph RAG')
add_normal(doc, 'Kiến trúc Graph RAG của hệ thống được chia thành hai giai đoạn chính:')
add_normal(doc, 'Giai đoạn 1 – Xây dựng đồ thị: Tài liệu được phân mảnh theo cấu trúc và bảng biểu, mỗi đoạn khởi tạo thành nút độc lập. Ba loại cạnh được thiết lập: cạnh cấu trúc (nội tuyến cùng tài liệu), cạnh siêu dữ liệu (cùng phòng ban), và cạnh ngữ nghĩa (cosine ≥ 0.7). Thuật toán Louvain phát hiện cộng đồng, tự động nhóm nút thành các phân vùng phòng ban.', first_indent=True)
add_normal(doc, 'Giai đoạn 2 – Truy xuất kết hợp đồ thị: Câu hỏi người dùng được mã hóa thành vector, định tuyến đến top-3 cộng đồng liên quan nhất. Trong mỗi cộng đồng, điểm tổng hợp từ Dense Vector Search và BM25 xác định top tài liệu ứng viên. BFS duyệt đồ thị từ các điểm neo để thu thêm nút lân cận có cosine ≥ 0.8.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 3.14 – Kiến trúc Graph RAG]', first_indent=False)

add_heading3(doc, '3.2.3. Thiết kế cơ sở dữ liệu')
add_normal(doc, '3.2.3.1. Cơ sở dữ liệu điểm sinh viên (PostgreSQL)', bold=True, first_indent=False)

db_table(doc, 'Bảng 3.12. Cấu trúc bảng Score – lưu trữ thông tin điểm số',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','id','INT','Khóa chính'],['2','student_code','VARCHAR(20)','Mã sinh viên'],
     ['3','subject_name','VARCHAR(100)','Tên môn học'],['4','semester','VARCHAR(50)','Học kỳ'],
     ['5','score_first','FLOAT4','Điểm thành phần 1'],['6','score_second','FLOAT4','Điểm thành phần 2'],
     ['7','score_final','FLOAT4','Điểm thi cuối kỳ'],['8','score_total','FLOAT4','Điểm tổng kết']],
    [1.0,3.5,3.5,6.5])

db_table(doc, 'Bảng 3.13. Cấu trúc bảng Structure – lưu trữ thông tin môn học',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','subject_id','INT','Khóa chính'],['2','subject_name','VARCHAR(100)','Tên môn học'],
     ['3','subject_credits','INT','Số tín chỉ môn học']],
    [1.0,3.5,3.5,6.5])

db_table(doc, 'Bảng 3.14. Cấu trúc bảng Student – lưu trữ thông tin sinh viên',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','student_code','INT','Mã sinh viên (khóa chính)'],['2','student_name','VARCHAR(100)','Tên sinh viên'],
     ['3','student_class','VARCHAR(20)','Lớp niên chế']],
    [1.0,3.5,3.5,6.5])

add_normal(doc, '3.2.3.2. Cơ sở dữ liệu lịch sử hội thoại (MongoDB)', bold=True, first_indent=False)
add_normal(doc, 'Hệ thống sử dụng MongoDB để lưu trữ dữ liệu linh hoạt như lịch sử hội thoại, tài khoản người dùng và cấu hình mô hình.')

db_table(doc, 'Bảng 3.15. Cấu trúc bảng Users – lưu trữ tài khoản và xác thực',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','user_id','String','ID người dùng (ObjectId)'],['2','username','String','Tên đăng nhập'],
     ['3','email','String','Địa chỉ email'],['4','password_hash','String','Mật khẩu đã mã hóa (bcrypt)'],
     ['5','role','String','Vai trò: user / admin / staff'],['6','student_code','String','Mã sinh viên (nếu có)'],
     ['7','created_at','DateTime','Thời gian tạo tài khoản'],['8','is_active','Boolean','Trạng thái tài khoản']],
    [1.0,3.5,3.0,7.0])

db_table(doc, 'Bảng 3.16. Cấu trúc bảng Conversations – lưu trữ cuộc trò chuyện',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','conversation_id','String','ID cuộc trò chuyện'],['2','user_id','String','ID người dùng sở hữu'],
     ['3','title','String','Tiêu đề cuộc trò chuyện'],['4','created_at','DateTime','Thời gian tạo'],
     ['5','updated_at','DateTime','Thời gian cập nhật lần cuối']],
    [1.0,3.5,3.0,7.0])

db_table(doc, 'Bảng 3.17. Cấu trúc bảng Messages – lưu trữ tin nhắn',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','message_id','String','ID tin nhắn'],['2','conversation_id','String','ID cuộc trò chuyện liên kết'],
     ['3','sender','String','Người gửi (user / assistant)'],['4','content','String','Nội dung tin nhắn'],
     ['5','created_at','DateTime','Thời gian gửi'],['6','attachments','Array','Danh sách file đính kèm'],
     ['7','metadata','Object','Thông tin model, số token, v.v.']],
    [1.0,3.5,3.0,7.0])

db_table(doc, 'Bảng 3.18. Cấu trúc bảng Rate_limits – kiểm soát tốc độ yêu cầu',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','user_id','String','ID người dùng'],['2','request_count','Int','Số yêu cầu đã gửi'],
     ['3','reset_at','DateTime','Thời gian đặt lại bộ đếm'],['4','is_blocked','Boolean','Trạng thái bị chặn']],
    [1.0,3.5,3.0,7.0])

db_table(doc, 'Bảng 3.19. Cấu trúc bảng LLM_models – quản lý mô hình AI',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','name','String','Tên mô hình AI'],['2','modelType','String','Loại: ollama / gemini / huggingface'],
     ['3','modelId','String','ID của mô hình'],['4','isActive','Boolean','Trạng thái đang hoạt động']],
    [1.0,3.5,3.0,7.0])

add_normal(doc, '3.2.3.3. Cơ sở dữ liệu vector embedding (Qdrant)', bold=True, first_indent=False)
add_normal(doc, 'Các văn bản hành chính của Học viện được lưu dưới dạng vector embedding trong Qdrant để phục vụ tìm kiếm ngữ nghĩa.')

db_table(doc, 'Bảng 3.20. Cấu trúc bảng User_document_embeddings',
    ['STT', 'Tên trường', 'Kiểu dữ liệu', 'Ý nghĩa'],
    [['1','file_id','VARCHAR','ID tệp gốc'],['2','chunk_index','INT','Chỉ số đoạn văn bản'],
     ['3','text','VARCHAR','Nội dung văn bản gốc của đoạn chunk'],
     ['4','embedding','VECTOR','Vector nhúng chiều cao'],
     ['5','metadata','JSON','Thông tin bổ sung: department, doc_type, section_title']],
    [1.0,3.5,3.0,7.0])

add_heading3(doc, '3.2.4. Thiết kế API Endpoints')
add_normal(doc, 'Backend cung cấp RESTful API được chia thành bốn nhóm chính: Chat, Xác thực, Quản trị và Dữ liệu RAG.')
add_normal(doc, '3.2.4.1. Nhóm API Chat', bold=True, first_indent=False)

api_table(doc, 'Bảng 3.21. API tra cứu điểm sinh viên',
    '/api/chat/score', 'POST',
    'student_code (String): Mã sinh viên; semester (String, tùy chọn): Học kỳ cần tra cứu',
    'Danh sách điểm theo học kỳ, bao gồm điểm thành phần và điểm tổng kết')

api_table(doc, 'Bảng 3.22. API tạo cuộc trò chuyện mới',
    '/api/chat/conversations', 'POST',
    'title (String): Tiêu đề cuộc trò chuyện',
    'Thông tin cuộc trò chuyện vừa tạo (conversation_id, title, created_at)')

api_table(doc, 'Bảng 3.23. API xoá cuộc trò chuyện',
    '/api/chat/conversations/:conversationId', 'DELETE',
    'conversationId (String): ID cuộc trò chuyện cần xóa',
    'Thông báo xóa thành công')

api_table(doc, 'Bảng 3.24. API lấy lịch sử chat của cuộc trò chuyện',
    '/api/chat/messages/:conversationId', 'GET',
    'conversationId (String): ID; skip (Number): Phân trang; limit (Number): Số bản ghi',
    'Danh sách tin nhắn trong cuộc trò chuyện')

api_table(doc, 'Bảng 3.25. API gửi tin nhắn trong cuộc trò chuyện',
    '/api/chat/:conversationId/messages', 'POST',
    'conversationId (String): ID; content (String): Nội dung tin nhắn; attachments (Array, tùy chọn)',
    'Tin nhắn phản hồi dạng streaming (Server-Sent Events)')

add_normal(doc, '3.2.4.2. Nhóm API Xác thực', bold=True, first_indent=False)

api_table(doc, 'Bảng 3.26. API đăng nhập hệ thống',
    '/api/auth/login', 'POST',
    'username (String): Tên đăng nhập; password (String): Mật khẩu',
    'access_token (JWT), refresh_token, thông tin người dùng')

api_table(doc, 'Bảng 3.27. API lấy thông tin người dùng hiện tại',
    '/api/auth/me', 'GET',
    'Authorization header: Bearer access_token',
    'Thông tin chi tiết người dùng đang đăng nhập')

api_table(doc, 'Bảng 3.28. API làm mới token',
    '/api/auth/refresh', 'POST',
    'refresh_token (String): Refresh token cần làm mới',
    'access_token mới')

add_normal(doc, '3.2.4.3. Nhóm API Quản trị (Admin)', bold=True, first_indent=False)

api_table(doc, 'Bảng 3.29. API tạo tài khoản người dùng',
    '/api/users/', 'POST',
    'username, password, email, role (String)',
    'Thông tin tài khoản vừa tạo')

api_table(doc, 'Bảng 3.30. API xem danh sách người dùng',
    '/api/users/', 'GET',
    'Authorization header Admin; skip, limit (phân trang)',
    'Danh sách tài khoản người dùng')

api_table(doc, 'Bảng 3.31. API cập nhật người dùng',
    '/api/users/admin/:user_id', 'PUT',
    'user_id (String): ID người dùng; studentCode, role, is_active (tùy chọn)',
    'Thông tin người dùng sau khi cập nhật')

api_table(doc, 'Bảng 3.32. API xoá người dùng',
    '/api/users/admin/:user_id', 'DELETE',
    'user_id (String): ID người dùng cần xóa',
    'Thông báo xóa thành công')

api_table(doc, 'Bảng 3.33. API xem model đang hoạt động',
    '/api/models/active', 'GET',
    'Authorization header (Bearer token)',
    'Thông tin mô hình AI đang được kích hoạt')

api_table(doc, 'Bảng 3.34. API kích hoạt model',
    '/api/models/activate/:model_id', 'POST',
    'model_id (String): ID model cần kích hoạt',
    'Xác nhận model đã được kích hoạt')

api_table(doc, 'Bảng 3.35. API tải lên model mới',
    '/api/models/upload', 'POST',
    'name, modelType, modelId, config (Object): Cấu hình mô hình',
    'Thông tin mô hình vừa được thêm vào hệ thống')

api_table(doc, 'Bảng 3.36. API lấy cấu hình giới hạn tốc độ',
    '/api/rate-limits/config', 'GET',
    'Authorization header Admin',
    'Cấu hình rate limit hiện tại (requests/phút, token/ngày, v.v.)')

api_table(doc, 'Bảng 3.37. API tạo/cập nhật cấu hình giới hạn tốc độ',
    '/api/rate-limits/config', 'POST/PUT',
    'max_requests_per_minute, max_tokens_per_day, max_file_size_mb (Number)',
    'Cấu hình giới hạn sau khi cập nhật')

api_table(doc, 'Bảng 3.38. API xem thống kê rate limit tất cả người dùng',
    '/api/rate-limits/stats', 'GET',
    'Authorization header Admin; page, limit (phân trang)',
    'Danh sách thống kê rate limit theo từng người dùng')

api_table(doc, 'Bảng 3.39. API xem tổng hợp sử dụng rate limit',
    '/api/rate-limits/summary', 'GET',
    'Authorization header Admin',
    'Tổng hợp: tổng request, tổng token đã dùng theo ngày/tháng')

api_table(doc, 'Bảng 3.40. API xem thống kê dashboard',
    '/api/stats/dashboard', 'GET',
    'Authorization header Admin',
    'Tổng số người dùng, tin nhắn, token sử dụng, trạng thái hệ thống')

api_table(doc, 'Bảng 3.41. API xem thống kê cuộc hội thoại',
    '/api/stats/conversations', 'GET',
    'Authorization header Admin; from_date, to_date (tùy chọn)',
    'Thống kê số lượng cuộc hội thoại và tin nhắn theo thời gian')

add_normal(doc, '3.2.4.4. Nhóm API Dữ liệu RAG', bold=True, first_indent=False)

api_table(doc, 'Bảng 3.42. API tải file huấn luyện RAG',
    '/api/rag/files/upload', 'POST',
    'file (multipart): File PDF/DOCX/TXT; department (String): Tên phòng ban',
    'Thông tin file vừa tải lên và trạng thái xử lý')

api_table(doc, 'Bảng 3.43. API liệt kê file huấn luyện RAG',
    '/api/rag/files', 'GET',
    'department (String, tùy chọn): Lọc theo phòng ban',
    'Danh sách file đã tải lên kèm số lượng chunk và trạng thái embedding')

api_table(doc, 'Bảng 3.44. API xoá file huấn luyện RAG',
    '/api/rag/files/:file_id', 'DELETE',
    'file_id (String): ID file cần xóa',
    'Thông báo xóa thành công, xóa cả embedding liên quan')

api_table(doc, 'Bảng 3.45. API rebuild toàn bộ RAG index',
    '/api/rag/rebuild', 'POST',
    'Authorization header Admin',
    'Xác nhận bắt đầu rebuild, trả về job_id để theo dõi tiến trình')

api_table(doc, 'Bảng 3.46. API rebuild RAG index theo phòng ban',
    '/api/rag/rebuild/:department', 'POST',
    'department (String): Tên phòng ban cần rebuild',
    'Xác nhận bắt đầu rebuild cho phòng ban cụ thể')

api_table(doc, 'Bảng 3.47. API liệt kê phòng ban dữ liệu RAG',
    '/api/rag/departments', 'GET',
    'Authorization header Admin',
    'Danh sách phòng ban kèm số lượng file và tài liệu')

api_table(doc, 'Bảng 3.48. API tạo thư mục dữ liệu RAG',
    '/api/rag/departments', 'POST',
    'name (String): Tên thư mục/phòng ban mới',
    'Thông tin thư mục vừa tạo')

api_table(doc, 'Bảng 3.49. API xoá thư mục dữ liệu RAG',
    '/api/rag/departments/:dept_name', 'DELETE',
    'dept_name (String): Tên phòng ban cần xóa',
    'Thông báo xóa thành công (bao gồm toàn bộ file và embedding trong thư mục)')

api_table(doc, 'Bảng 3.50. API tải file từ thư mục dữ liệu RAG',
    '/api/rag/departments/:dept_name/files', 'POST',
    'dept_name (String): Tên phòng ban; file (multipart): File cần tải lên',
    'Thông tin file vừa tải lên và trạng thái xử lý embedding')

add_heading3(doc, '3.2.5. Thiết kế giao diện người dùng')
add_normal(doc, 'Giao diện người dùng được thiết kế theo hai nhóm chính: giao diện dành cho người dùng thông thường và giao diện quản trị (Admin Dashboard).')

gd_items = [
    ('3.15', 'trang chủ', 'Trang chủ giới thiệu hệ thống, mô tả khả năng của trợ lý ảo ACTVN-AGENT và hướng dẫn truy cập.'),
    ('3.16', 'hội thoại chính', 'Giao diện chat tương tự ứng dụng nhắn tin, tích hợp thanh điều hướng lịch sử hội thoại bên trái và khung nhập liệu bên dưới.'),
    ('3.17', 'tra cứu điểm', 'Hiển thị kết quả điểm số dưới dạng bảng có cấu trúc ngay trong luồng hội thoại.'),
    ('3.18', 'truy vấn tài liệu hành chính', 'Câu trả lời hiển thị định dạng văn bản có cấu trúc, kèm trích dẫn nguồn tài liệu.'),
    ('3.19', 'hỏi đáp với tệp tin', 'Người dùng tải lên tài liệu cá nhân (PDF/DOCX) và đặt câu hỏi về nội dung tệp.'),
    ('3.20', 'quản lý tệp tin tải lên', 'Danh sách file lưu trữ, số vector embedding đã tạo và thao tác xóa file.'),
    ('3.21', 'Admin Dashboard', 'Biểu đồ thống kê thời gian thực: tổng request, token sử dụng, danh sách hoạt động gần đây.'),
    ('3.22', 'quản lý người dùng', 'Bảng theo dõi tài khoản: trạng thái, vai trò, đăng nhập lần cuối và thao tác quản lý.'),
    ('3.23', 'quản lý lịch sử hội thoại', 'Liệt kê chi tiết các phiên thảo luận của tất cả người dùng kèm số token tiêu tốn.'),
    ('3.24', 'giới hạn sử dụng', 'Cấu hình số yêu cầu tối đa và lượng token theo từng khung thời gian cho các nhóm người dùng.'),
    ('3.25', 'quản lý mô hình LLM', 'Hiển thị mô hình đang hoạt động và danh sách mô hình có sẵn (Gemini, Qwen, Llama).'),
    ('3.26', 'giới hạn tệp tin tải lên', 'Thiết lập định mức số lượng và dung lượng tệp tin cho Admin và người dùng.'),
    ('3.27', 'quản lý dữ liệu RAG', 'Tải lên tài liệu huấn luyện theo thư mục phòng ban, theo dõi trạng thái và rebuild index.'),
    ]

for fig_num, name, desc in gd_items:
    add_normal(doc, f'Giao diện {name}: {desc}')
    add_normal(doc, f'[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình {fig_num} – Giao diện {name}]', first_indent=False)

# 3.3
add_heading2(doc, '3.3. Xây dựng và Cài đặt phần mềm')
add_heading3(doc, '3.3.1. Công cụ và công nghệ sử dụng')
add_normal(doc, 'Hệ thống sử dụng Python làm ngôn ngữ phát triển chính cho backend, phần dịch vụ API được xây dựng bằng FastAPI theo mô hình bất đồng bộ (async/await). Các công nghệ chính:')
add_bullet(doc, 'FastAPI (Python): Framework API bất đồng bộ, hỗ trợ tự động sinh tài liệu OpenAPI/Swagger.')
add_bullet(doc, 'LangChain: Chuẩn hóa quy trình gọi LLM, quản lý prompt và kết nối các thành phần RAG.')
add_bullet(doc, 'FAISS/Qdrant: Lưu trữ và tìm kiếm vector embedding tốc độ cao.')
add_bullet(doc, 'ReactJS: Giao diện web người dùng, giao tiếp với Backend qua API JWT.')
add_bullet(doc, 'Docker + Docker Compose: Container hóa và quản lý triển khai.')
add_bullet(doc, 'MongoDB: Lưu lịch sử hội thoại và cấu hình mô hình linh hoạt.')
add_bullet(doc, 'PostgreSQL: Lưu trữ dữ liệu điểm số và thông tin sinh viên.')
add_bullet(doc, 'nomic-embed-text: Mô hình embedding tiếng Việt triển khai qua Ollama.')

add_heading3(doc, '3.3.2. LLM Factory Pattern và Token Management')
add_normal(doc, 'Để hỗ trợ nhiều nhà cung cấp mô hình và dễ dàng thay đổi cấu hình, hệ thống áp dụng LLM Factory Pattern – một lớp trừu tượng chuẩn hóa giao tiếp với mọi nhà cung cấp (Gemini API, Ollama, HuggingFace). Lợi ích: giảm phụ thuộc trực tiếp, chuẩn hóa luồng gọi mô hình và đơn giản hóa tích hợp mô hình mới.')
add_normal(doc, 'Token Management theo dõi số lượng token đầu vào/đầu ra của từng yêu cầu, thiết lập hạn mức theo tài khoản/nhóm người dùng, ghi log phục vụ thống kê và kết hợp với Rate Limiting để ngăn ngừa lạm dụng tài nguyên.')

add_heading3(doc, '3.3.3. Backend API và Authentication')
add_normal(doc, 'Backend triển khai cơ chế bảo mật JWT (JSON Web Token) với Access Token thời hạn ngắn và Refresh Token thời hạn dài. Sau đăng nhập, mọi request cần mang Bearer token trong Authorization header. Rate Limiting được áp dụng theo endpoint và nhóm người dùng để bảo vệ hệ thống trước tấn công DDoS và lạm dụng API.')

add_heading3(doc, '3.3.4. Web App và Mobile App')
add_normal(doc, 'Web App ReactJS đã được hoàn thiện với đầy đủ chức năng cho cả người dùng và Admin. Ứng dụng giao tiếp với Backend qua API JWT, hỗ trợ streaming response cho chức năng chat AI. Mobile App (React Native) đã được xác định trong kiến trúc tổng thể và là hướng phát triển tiếp theo, sử dụng lại toàn bộ API đã xây dựng.')

add_heading2(doc, '3.4. Tổng kết chương 3')
add_normal(doc, 'Chương 3 đã trình bày toàn bộ quá trình phân tích, thiết kế và cài đặt hệ thống Chatbot thông minh tại Học viện Kỹ thuật Mật mã. Từ việc xác định rõ yêu cầu qua 11 use case, thiết kế kiến trúc Graph RAG chuyên biệt, xây dựng cơ sở dữ liệu đa tầng (PostgreSQL + MongoDB + Qdrant), đến thiết kế 50 API endpoint đầy đủ và giao diện người dùng thân thiện – hệ thống đã được hiện thực hóa đầy đủ. Kết quả là một nền tảng Chatbot hoàn chỉnh sẵn sàng cho giai đoạn thực nghiệm và đánh giá ở chương tiếp theo.')

doc.save('document/TAI_LIEU_SAN_PHAM_CHUAN_FORMAT_part3.docx')
print("Saved part3 – Chuong 3 done.")
