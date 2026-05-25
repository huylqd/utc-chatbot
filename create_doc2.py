# -*- coding: utf-8 -*-
"""Phase 2: Chương 1 + Chương 2"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# re-use same helpers
def fmt_run(run, bold=False, italic=False, size=13):
    run.font.name = 'Times New Roman'; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic

def add_normal(doc, text, first_indent=True, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); p.style = doc.styles['Normal']
    p.paragraph_format.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.first_line_indent = Cm(1.0) if first_indent else Cm(0)
    run = p.add_run(text); fmt_run(run, bold=bold); return p

def add_heading1(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 1']
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(18); run.font.bold=True
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 2']
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(14); run.font.bold=True
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles['Heading 3']
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(13); run.font.bold=True
    return p

def add_bullet(doc, text, indent_level=1, bold=False):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Cm(indent_level * 0.75)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(f'• {text}'); fmt_run(run, bold=bold); return p

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text); run.font.name='Times New Roman'; run.font.size=Pt(12); run.font.italic=True
    return p

def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); run._r.append(br)

def set_cell_bg(cell, color='BDD7EE'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color); shd.set(qn('w:color'),'auto'); shd.set(qn('w:val'),'clear')
    tcPr.append(shd)

def make_table(doc, headers, rows, col_widths=None, caption=None):
    if caption:
        add_caption(doc, caption)
    n_cols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hrow.cells[i], 'BDD7EE')
        p = hrow.cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.font.name='Times New Roman'; run.font.size=Pt(11); run.font.bold=True
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri+1]
        for ci, val in enumerate(row_data):
            p = row.cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci==0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val)); run.font.name='Times New Roman'; run.font.size=Pt(11)
    if col_widths:
        for ci,w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

# ── Load part1 ────────────────────────────────────────────────────────────────
doc = Document('document/TAI_LIEU_SAN_PHAM_CHUAN_FORMAT_part1.docx')

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 1
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI VÀ CƠ SỞ LÝ THUYẾT')

# 1.1
add_heading2(doc, '1.1. Khảo sát hệ thống Chatbot và bài toán hỏi đáp hiện có')
add_heading3(doc, '1.1.1. Các hệ thống Chatbot giáo dục trên thế giới và Việt Nam')
add_normal(doc, 'Trong giai đoạn 2020–2026, lĩnh vực trí tuệ nhân tạo ứng dụng trong giáo dục đã trải qua một quá trình chuyển biến căn bản. Sự xuất hiện của các mô hình ngôn ngữ lớn như GPT-4, Gemini và Claude đã mở ra thế hệ Chatbot giáo dục mới – không chỉ trả lời câu hỏi mà còn có khả năng hỗ trợ cá nhân hóa lộ trình học tập, đánh giá phản hồi và tương tác đa lượt.')
add_normal(doc, 'Ba xu hướng chủ đạo đang định hình thế hệ Chatbot giáo dục hiện tại: (1) cá nhân hóa lộ trình học tập; (2) tích hợp đa phương thức (văn bản, hình ảnh, giọng nói); và (3) hỗ trợ suy luận Socratic thay vì cung cấp đáp án trực tiếp.')

add_heading3(doc, '1.1.2. Các hệ thống Chatbot giáo dục tiêu biểu trên thế giới')
add_normal(doc, 'Khanmigo (Khan Academy, 2023) được xây dựng trên nền GPT-4 kết hợp kho học liệu sư phạm của Khan Academy. Hệ thống áp dụng triết lý Socratic – dẫn dắt người học tự khám phá thay vì cung cấp đáp án trực tiếp. Ngoài vai trò gia sư, Khanmigo còn hỗ trợ giáo viên thiết kế bài giảng và tạo rubric đánh giá.')
add_normal(doc, 'Duolingo Max (2023) là gói dịch vụ cao cấp tích hợp GPT-4 cho nền tảng học ngoại ngữ Duolingo, nổi bật với tính năng Explain My Answer phân tích lý do đúng/sai, và tính năng giải thích phát âm chuyên biệt theo từng ngôn ngữ. Q-Chat (Quizlet, 2023) tích hợp thuật toán lặp lại ngắt quãng vào luồng hội thoại giúp tối ưu việc ghi nhớ kiến thức.')
add_normal(doc, 'Tại Việt Nam, ViGPT (VinBigData, 2023–2024) là mô hình ngôn ngữ lớn được huấn luyện chuyên biệt cho tiếng Việt, đánh dấu bước tiến quan trọng trong phát triển AI giáo dục bản địa. ELSA Speak chuyên biệt hóa về phát âm tiếng Anh với công nghệ nhận diện giọng nói tiên tiến.')

# Table 1.1
make_table(doc,
    ['Hệ thống', 'Năm', 'Công nghệ lõi', 'Triết lý sư phạm', 'Hạn chế chính'],
    [
        ['Khanmigo', '2023', 'GPT-4, RAG, kho học liệu Khan', 'Socratic – dẫn dắt tự khám phá', 'Chi phí cao; chỉ tiếng Anh'],
        ['Duolingo Max', '2023', 'GPT-4 + ASR chuyên biệt', 'Thực hành giao tiếp tình huống thực', 'Chỉ dùng học ngoại ngữ'],
        ['Q-Chat', '2023', 'GPT-4 + SRS (Ebbinghaus)', 'Lặp lại ngắt quãng thông minh', 'Phụ thuộc thẻ ghi nhớ sẵn có'],
        ['Socratic', '2019/2022', 'Computer Vision, OCR, LLM', 'Tra cứu đa nguồn minh bạch', 'Không hỗ trợ hội thoại đa lượt'],
        ['MATHia', '2019–nay', 'Cognitive Tutor, ML', 'Mô hình hóa nhận thức từng bước', 'Chỉ môn Toán'],
        ['ViGPT', '2023–2024', 'LLM tiếng Việt chuyên biệt', 'Hỗ trợ ngữ cảnh Việt Nam', 'Đang phát triển, chưa hoàn thiện'],
    ],
    col_widths=[2.5, 1.5, 3.5, 3.5, 3.0],
    caption='Bảng 1.1. So sánh các hệ thống Chatbot giáo dục tiêu biểu'
)

add_heading3(doc, '1.1.3. Đánh giá ưu nhược điểm và định hướng phát triển')
add_normal(doc, 'Để đưa ra lựa chọn kiến trúc có căn cứ, bảng dưới đây đánh giá bốn hướng tiếp cận phổ biến theo các tiêu chí then chốt.')

make_table(doc,
    ['Tiêu chí đánh giá', 'Rule-based', 'Traditional ML/NLP', 'LLM thuần túy', 'RAG + LLM (Đề tài)'],
    [
        ['Hiểu ngôn ngữ tự nhiên', 'Thấp', 'Trung bình', 'Cao', 'Cao'],
        ['Xử lý tài liệu nội bộ', 'Hạn chế', 'Hạn chế', 'Không cập nhật được', '✔ Linh hoạt, thời gian thực'],
        ['Kiểm soát hallucination', 'Cao (cứng nhắc)', 'Trung bình', 'Thấp (3–8%)', '✔ Cao (1–2%)'],
        ['Truy ngược nguồn gốc', 'Không', 'Không', 'Không', '✔ Có (trích dẫn tài liệu)'],
        ['Khả năng mở rộng', 'Rất thấp', 'Thấp', 'Trung bình (fine-tune)', '✔ Cao (cập nhật vector DB)'],
    ],
    col_widths=[4.0, 2.5, 3.0, 3.0, 3.5],
    caption='Bảng 1.2. So sánh các hướng tiếp cận xây dựng hệ thống Chatbot hỏi đáp tài liệu'
)

add_normal(doc, 'Kết quả phân tích cho thấy kiến trúc RAG kết hợp LLM vượt trội ở hầu hết các tiêu chí then chốt. Đây là lý do đề tài lựa chọn kiến trúc này làm nền tảng. Điểm mạnh cốt lõi gồm: tính chính xác có thể xác minh nguồn gốc, khả năng cập nhật kiến thức linh hoạt không cần huấn luyện lại mô hình, và đáp ứng yêu cầu bảo mật bằng triển khai on-premise.')

# 1.2
add_heading2(doc, '1.2. Tìm hiểu về Large Language Models và Small Language Models')
add_heading3(doc, '1.2.1. Khái niệm và đặc điểm của LLM')
add_normal(doc, 'Mô hình ngôn ngữ lớn (Large Language Model – LLM) là một lớp mô hình học sâu được huấn luyện trên khối lượng dữ liệu văn bản khổng lồ với mục tiêu học cách dự đoán và sinh ra ngôn ngữ tự nhiên. Mọi LLM hiện đại đều được xây dựng trên kiến trúc Transformer, với cơ chế Self-Attention cho phép mô hình "nhìn thấy" toàn bộ ngữ cảnh trong một bước xử lý.')
add_normal(doc, 'LLM được huấn luyện theo quy trình hai giai đoạn: Pre-training trên tập dữ liệu khổng lồ (Wikipedia, sách, mã nguồn) và Fine-tuning trên dữ liệu chuyên biệt theo tác vụ. Một đặc điểm nổi bật là hiện tượng emergent abilities – các năng lực xuất hiện đột ngột khi mô hình đạt ngưỡng kích thước nhất định, ví dụ như suy luận chuỗi suy nghĩ (chain-of-thought).')

add_heading3(doc, '1.2.2. Sự ra đời của mô hình ngôn ngữ nhỏ (SLM)')
add_normal(doc, 'Mô hình ngôn ngữ nhỏ (Small Language Model – SLM) thường được định nghĩa là các mô hình dưới 10 tỷ tham số. Sự dịch chuyển sang SLM không phải là bước lùi về công nghệ mà là phản ứng thực tế: chi phí triển khai thấp hơn nhiều, có thể chạy trên phần cứng phổ thông, độ trễ thấp hơn và dễ kiểm soát bảo mật dữ liệu.')

make_table(doc,
    ['Tiêu chí', 'LLM (>50B tham số)', 'SLM (<10B tham số)'],
    [
        ['Đại diện tiêu biểu', 'GPT-4, Claude 3 Opus, Gemini Ultra', 'Phi-3 Mini, Mistral 7B, Gemma 7B, LLaMA-3 8B'],
        ['Chi phí API/vận hành', '30–60 USD/M tokens', '<1 USD/M tokens (tự triển khai)'],
        ['Yêu cầu phần cứng', 'Cụm GPU A100/H100 (>$100,000)', '1× GPU RTX 4090 hoặc A10 (~$2,000–10,000)'],
        ['Thời gian phản hồi', '500ms – 5 giây', '100–500ms (on-premise, ổn định)'],
        ['Bảo mật dữ liệu', 'Dữ liệu gửi lên cloud', 'Triển khai hoàn toàn nội bộ (on-premise)'],
        ['Khả năng fine-tune', 'Khó, tốn kém', 'Dễ dàng với LoRA/QLoRA'],
    ],
    col_widths=[4.5, 5.0, 5.0],
    caption='Bảng 1.3. So sánh toàn diện LLM và SLM'
)

add_heading3(doc, '1.2.3. Các mô hình LLM/SLM phổ biến')
add_normal(doc, 'Google Gemini là dòng mô hình đa phương thức (văn bản, ảnh, âm thanh, video, PDF) do Google DeepMind phát triển, ra mắt tháng 12/2023. Gemini 2.5 Pro với cửa sổ ngữ cảnh 1 triệu token là lợi thế cạnh tranh lớn nhất, đặc biệt phù hợp với các bài toán xử lý tài liệu lớn.')
add_normal(doc, 'Meta LLaMA là dòng mô hình mã nguồn mở với giấy phép thương mại. LLaMA 4 (tháng 4/2025) với kiến trúc MoE multimodal là bước nhảy vọt lớn nhất của dòng này. Alibaba Qwen có ưu thế đặc biệt về xử lý tiếng Việt và các ngôn ngữ Đông Nam Á, được cấp phép Apache 2.0 cho hầu hết phiên bản.')

make_table(doc,
    ['Tiêu chí', 'Gemini 2.5 Pro', 'LLaMA 4 Maverick', 'Qwen 2.5/Qwen 3'],
    [
        ['Nhà phát triển', 'Google DeepMind', 'Meta AI', 'Alibaba Cloud'],
        ['Context window', '1M token', '1M token (Maverick)', '128K token (Q2.5); 32K–128K (Q3)'],
        ['Đa phương thức', 'Văn bản, ảnh, âm thanh, video, PDF', 'Văn bản + hình ảnh', 'Văn bản + ảnh + audio (Omni)'],
        ['Tiếng Việt', 'Tốt', 'Trung bình', 'Tốt (benchmark chính thức)'],
        ['Giấy phép', 'API thương mại', 'Mã nguồn mở', 'Apache 2.0 (hầu hết phiên bản)'],
        ['Triển khai on-premise', 'Không', 'Có', 'Có'],
    ],
    col_widths=[4.0, 4.0, 3.5, 3.5],
    caption='Bảng 1.4. So sánh chi tiết các mô hình LLM/SLM tiêu biểu (cập nhật Q1/2026)'
)

# 1.3
add_heading2(doc, '1.3. Kỹ thuật Prompt Engineering')
add_heading3(doc, '1.3.1. Khái niệm và tầm quan trọng')
add_normal(doc, 'Prompt Engineering là quá trình thiết kế, tối ưu hóa và tinh chỉnh các đầu vào văn bản nhằm điều khiển hành vi của mô hình ngôn ngữ lớn để tạo ra đầu ra mong muốn. Cùng một LLM, prompt được thiết kế kỹ có thể cải thiện độ chính xác từ 45% lên 88% trên các tác vụ phân loại văn bản chuyên ngành.')

make_table(doc,
    ['Thành phần', 'Mô tả', 'Ví dụ'],
    [
        ['Instruction', 'Chỉ thị tác vụ cụ thể', '"Tóm tắt văn bản trong 3 câu"'],
        ['Context', 'Ngữ cảnh và thông tin nền', 'Role, lịch sử hội thoại, chính sách'],
        ['Input Data', 'Dữ liệu cần xử lý', 'Văn bản, câu hỏi, tài liệu'],
        ['Examples', 'Ví dụ minh họa', 'Cặp input → output mẫu'],
        ['Output Format', 'Định dạng đầu ra mong muốn', 'JSON, bullet points, 200 từ'],
    ],
    col_widths=[3.5, 5.0, 6.0],
    caption='Bảng 1.5. Bảng thành phần cấu tạo của một prompt'
)

add_heading3(doc, '1.3.2. Các kỹ thuật cơ bản')
add_normal(doc, 'Zero-shot Prompting yêu cầu mô hình thực hiện tác vụ chỉ với mô tả chỉ thị, không có ví dụ. Few-shot Prompting cung cấp một số cặp input–output mẫu, đặc biệt hiệu quả với tác vụ yêu cầu định dạng đầu ra đặc thù. ReAct (Reasoning + Acting) kết hợp suy luận với gọi công cụ bên ngoài, cho phép mô hình lý luận từng bước và quyết định hành động phù hợp.')

# 1.4
add_heading2(doc, '1.4. Tìm hiểu về Retrieval-Augmented Generation (RAG)')
add_heading3(doc, '1.4.1. Tổng quan về RAG')
add_normal(doc, 'Retrieval-Augmented Generation (RAG) là kiến trúc tiên tiến kết hợp giữa khả năng truy xuất thông tin và sinh ngôn ngữ tự nhiên. Ý tưởng cốt lõi: thay vì chỉ dựa vào kiến thức được mã hóa trong tham số mô hình, RAG cho phép hệ thống chủ động tìm kiếm thông tin liên quan từ kho tài liệu bên ngoài và đưa vào ngữ cảnh trả lời.')
add_normal(doc, 'Một hệ thống RAG hoàn chỉnh vận hành qua ba luồng chính: (1) Indexing – tiền xử lý và lập chỉ mục tài liệu; (2) Retrieval – truy xuất thông tin liên quan dựa trên truy vấn; (3) Generation – sinh câu trả lời dựa trên ngữ cảnh được truy xuất.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 1.1 – Kiến trúc tổng quan hệ thống RAG]', first_indent=False, bold=False)

add_heading3(doc, '1.4.2. Vector RAG truyền thống')
add_normal(doc, 'Vector RAG lưu trữ tài liệu dưới dạng vector embedding trong cơ sở dữ liệu vector (Pinecone, Weaviate, Milvus, ChromaDB, FAISS). Khi có truy vấn, hệ thống tính độ tương đồng Cosine giữa vector câu hỏi và vector tài liệu để truy xuất những đoạn văn bản liên quan nhất.')
add_normal(doc, 'Chiến lược Document Chunking đóng vai trò then chốt. Các phương pháp nâng cao bao gồm Semantic Chunking (chia theo ranh giới ngữ nghĩa), Hierarchical Chunking (cấu trúc cây từ tóm tắt đến chi tiết), và Parent-Child Chunking (chunk nhỏ để tìm kiếm, chunk cha để sinh ngữ cảnh đầy đủ).')

add_heading3(doc, '1.4.3. Hạn chế của Vector RAG và Graph RAG')
add_normal(doc, 'Vector RAG truyền thống bộc lộ ba hạn chế chính khi xử lý tài liệu hành chính phức tạp: (1) độ nhạy với từ khóa – thường bỏ sót tài liệu liên quan có cách diễn đạt khác; (2) phân mảnh ngữ cảnh – bước chunking có thể cắt đứt luồng thông tin logic; (3) thiếu khả năng nhận thức mối liên kết giữa các điều khoản, quy định liên quan.')
add_normal(doc, 'Graph RAG ra đời để khắc phục những hạn chế này bằng cách biểu diễn tri thức dưới dạng đồ thị, cho phép truy xuất thông tin qua duyệt đồ thị (graph traversal) thay vì chỉ tính độ tương đồng vector. Kiến trúc Graph RAG bao gồm bốn giai đoạn: Document to Graph Conversion, Entity và Relationship Extraction, Graph Construction, và Graph Traversal for Retrieval.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 1.2 – Pipeline xử lý tổng thể của Graph RAG]', first_indent=False)
add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 1.3 – Knowledge Graph kết nối thực thể qua quan hệ ngữ nghĩa]', first_indent=False)

add_heading3(doc, '1.4.4. Các kỹ thuật nâng cao trong RAG')
add_normal(doc, 'Hybrid Search kết hợp tìm kiếm thưa thớt (BM25 – từ khóa) và tìm kiếm dày đặc (Dense Vector), giúp bắt cả từ khóa chính xác lẫn ngữ nghĩa tổng quát. Re-ranking áp dụng mô hình Cross-Encoder để xếp hạng lại kết quả truy xuất ban đầu theo độ liên quan thực sự.')
add_normal(doc, 'Agentic RAG là hướng tiếp cận tiên tiến nhất, nơi LLM hoạt động như một tác nhân có khả năng lập kế hoạch, phân rã câu hỏi phức tạp, gọi nhiều công cụ tìm kiếm và tổng hợp kết quả theo nhiều bước.')

# 1.5
add_heading2(doc, '1.5. Tìm hiểu về Knowledge Graph')
add_heading3(doc, '1.5.1. Khái niệm và cấu trúc')
add_normal(doc, 'Knowledge Graph là cấu trúc dữ liệu dạng đồ thị để biểu diễn tri thức dưới dạng các thực thể (Entity) và mối quan hệ (Relation) giữa chúng. Đơn vị cơ bản là bộ ba (Triple): Subject – Predicate – Object. Ví dụ: "Sinh viên AT060001 – học tại – Học viện KT Mật mã".')
add_normal(doc, 'Về cấu trúc vật lý, Knowledge Graph gồm: Node (thực thể), Edge (quan hệ có hướng mang nhãn ngữ nghĩa), và Attribute (cặp khóa-giá trị gắn vào node/edge). Tầng Ontology định nghĩa Classes, Properties và Constraints để đảm bảo tính nhất quán.')

add_heading3(doc, '1.5.2. Graph Database và ứng dụng trong RAG')
add_normal(doc, 'Graph Database (ví dụ: Neo4j) là hệ thống lưu trữ tối ưu cho dữ liệu đồ thị, sử dụng ngôn ngữ truy vấn Cypher để mô tả pattern đồ thị. Khi kết hợp với RAG, Knowledge Graph giải quyết bài toán truy xuất thông tin có mối liên kết đa cấp, ví dụ: "Để điều kiện A áp dụng, cần thỏa mãn điều kiện B và C từ các quy chế khác nhau".')

# 1.6
add_heading2(doc, '1.6. Tìm hiểu về Agent và Multi-Agent Systems')
add_heading3(doc, '1.6.1. Khái niệm AI Agent')
add_normal(doc, 'AI Agent là hệ thống trí tuệ nhân tạo có khả năng tự chủ nhận thức môi trường, lập kế hoạch và thực hiện hành động để đạt mục tiêu. Một AI Agent điển hình gồm: Bộ nhận thức (thu thập đầu vào), Bộ nhớ (lưu ngữ cảnh ngắn/dài hạn), Lập kế hoạch (phân rã mục tiêu), Sử dụng công cụ (gọi API, tìm kiếm web, truy vấn DB) và Hành động (thực thi kết quả).')

add_heading3(doc, '1.6.2. Kiến trúc ReAct và Multi-Agent')
add_normal(doc, 'ReAct Pattern (Reasoning + Acting) cho phép Agent hoạt động theo vòng lặp động ba bước: Thought (phân tích và lập kế hoạch) → Action (gọi công cụ phù hợp) → Observation (nhận kết quả và quyết định bước tiếp). Vòng lặp kết thúc khi Agent đánh giá thông tin đã đủ để trả lời.')
add_normal(doc, 'Multi-Agent System (MAS) là kiến trúc nhiều Agent độc lập phối hợp giải quyết bài toán phức tạp. Trong đề tài này, mô hình Supervisor–Worker được áp dụng: Supervisor Agent trung tâm phân tích và định tuyến truy vấn đến các Worker Agent chuyên biệt (RAG Chatbot Agent, Score Lookup Agent, File Analysis Agent).')

# 1.7 Tổng kết
add_heading2(doc, '1.7. Tổng kết chương 1')
add_normal(doc, 'Chương 1 đã trình bày tổng quan về thực trạng các hệ thống Chatbot giáo dục và các cơ sở lý thuyết cốt lõi cho đề tài. Qua khảo sát, kiến trúc RAG kết hợp LLM được xác định là hướng tiếp cận phù hợp nhất cho bài toán hỏi đáp tài liệu nội bộ doanh nghiệp. Đặc biệt, Graph RAG vượt trội so với Vector RAG truyền thống trong xử lý tài liệu hành chính có cấu trúc phức tạp và mối liên kết đa chiều. Nền tảng lý thuyết về Knowledge Graph, ReAct Agent và Multi-Agent System sẽ là kim chỉ nam cho thiết kế hệ thống ở các chương tiếp theo.')

# ═══════════════════════════════════════════════════════════════════════════════
# CHƯƠNG 2
# ═══════════════════════════════════════════════════════════════════════════════
page_break(doc)
add_heading1(doc, 'CHƯƠNG 2. PHƯƠNG PHÁP VÀ KỸ THUẬT ĐỀ XUẤT')

add_heading2(doc, '2.1. Tổng quan phương pháp đề xuất')
add_heading3(doc, '2.1.1. Kiến trúc luồng xử lý tổng thể')
add_normal(doc, 'Để giải quyết những hạn chế của hệ thống tra cứu truyền thống và kiểm soát rủi ro "ảo giác" (hallucination) của LLM, hệ thống được thiết kế theo mô hình Hệ thống Đa tác tử (Multi-Agent System). Kiến trúc này xoay quanh một Supervisor Agent trung tâm chịu trách nhiệm phân tích, định tuyến và tổng hợp toàn bộ luồng xử lý.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 2.1 – Kiến trúc xử lý tổng thể]', first_indent=False)

add_normal(doc, 'Quá trình vận hành hệ thống được chia thành 4 giai đoạn cốt lõi:')
add_bullet(doc, 'Tiếp nhận và Phân tích: Supervisor Agent nhận truy vấn từ người dùng (văn bản hoặc file đính kèm), phân tích ý định và giải quyết đồng tham chiếu trong hội thoại đa lượt.')
add_bullet(doc, 'Phân luồng và Định tuyến: Dựa trên ý định, Supervisor định tuyến đến luồng phù hợp: (a) Luồng truy vấn tri thức qua Graph RAG; (b) Luồng tra cứu điểm số qua cơ sở dữ liệu cấu trúc; (c) Luồng bất đồng bộ cho xử lý tài liệu nặng.')
add_bullet(doc, 'Thực thi chuyên biệt: Mỗi tác tử chuyên biệt tự chủ gọi công cụ tương ứng (Vector Search, Graph Traversal, DB Query, File Parser) và thu thập thông tin.')
add_bullet(doc, 'Tổng hợp và Sinh phản hồi: Dữ liệu thô từ các tác tử được hội tụ về module sinh ngôn ngữ, kết hợp với ngữ cảnh hội thoại để tạo câu trả lời tự nhiên, chính xác.')

add_heading3(doc, '2.1.2. Luồng xử lý Chatbot truy vấn')
add_normal(doc, 'Khi truy vấn được định tuyến vào luồng giải đáp học vụ, ReAct Agent không sử dụng sinh văn bản một chiều mà vận hành theo vòng lặp Thought-Action-Observation. Tác tử kích hoạt cơ chế tìm kiếm kép: Vector Search để trích xuất chunks ngữ nghĩa gần nhất, và Graph Traversal để men theo các cạnh cấu trúc thu thập các điều khoản liên đới. Kết quả được đánh giá tính đủ; nếu chưa đủ, tác tử lập tức thực hiện thêm vòng truy xuất.')
add_normal(doc, 'Một tính năng quan trọng là cơ chế Bypass cơ sở dữ liệu: khi Supervisor phát hiện định danh sinh viên (ví dụ: mã AT... hoặc CT...) trong câu hỏi, luồng dữ liệu được định tuyến trực tiếp đến PostgreSQL để tra cứu điểm số theo thời gian thực, bỏ qua toàn bộ hệ thống RAG nhằm đảm bảo tính chính xác và giảm độ trễ.')

# 2.2
add_heading2(doc, '2.2. Thu thập và Tiền xử lý dữ liệu')
add_heading3(doc, '2.2.1. Thu thập tài liệu')
add_normal(doc, 'Hệ thống áp dụng chiến lược thu thập và phân loại dữ liệu theo cơ cấu tổ chức hành chính của Đại Học Giao Thông Vận Tải, bao gồm ba nhóm phòng ban chính:')
add_bullet(doc, 'Phòng Đào tạo: Khung chương trình đào tạo, chuẩn đầu ra tin học/ngoại ngữ, hệ thống tín chỉ và quy chế đào tạo đại học.')
add_bullet(doc, 'Phòng Khảo thí và Đảm bảo chất lượng: Quy chế tổ chức thi, cách tính điểm trung bình, quy định phúc khảo và kỷ luật thi.')
add_bullet(doc, 'Viện Nghiên cứu & Phòng Công tác sinh viên: Quy định nghiên cứu khoa học sinh viên, học bổng và các thủ tục hành chính nội bộ.')

add_heading3(doc, '2.2.2. Trích xuất nội dung và chuẩn hóa Markdown')
add_normal(doc, 'Dữ liệu đầu vào tồn tại dưới định dạng không có cấu trúc chuẩn như PDF, DOCX hoặc ảnh scan. Điểm đột phá của phương pháp là toàn bộ tài liệu sau khi bóc tách sẽ được tự động chuyển đổi thống nhất sang định dạng Markdown. Lý do:')
add_bullet(doc, 'Bảo toàn cấu trúc phân cấp đồ thị: Markdown duy trì hệ thống cây phân cấp qua các thẻ tiêu đề (H1/H2/H3), cho phép Graph RAG nhận biết đây là node "Chương" hay "Điều".')
add_bullet(doc, 'Giữ nguyên cấu trúc bảng biểu: Định dạng bảng Markdown bảo toàn ma trận hàng/cột, tránh phá vỡ quan hệ dữ liệu khi trích xuất.')

add_heading3(doc, '2.2.3. Chiến lược Chunking và trích xuất Metadata')
add_normal(doc, 'Các mô hình embedding và LLM có giới hạn cửa sổ ngữ cảnh, do đó tài liệu phải được chia nhỏ một cách thông minh. Hệ thống áp dụng hai chiến lược chuyên biệt:')
add_bullet(doc, 'Phân mảnh bảo toàn bảng biểu: Thuật toán tự động dò tìm ranh giới bảng Markdown và đảm bảo toàn bộ bảng nằm trong một chunk duy nhất.')
add_bullet(doc, 'Phân mảnh nhận thức cấu trúc pháp quy: Hệ thống được cấu hình Regular Expressions để nhận biết ranh giới điều khoản (Chương, Điều, Khoản, Điểm), đảm bảo ranh giới cắt nằm tại điểm chuyển giao điều khoản.')
add_normal(doc, 'Mỗi chunk sinh ra được gán đầy đủ metadata: source_file, department_owner, chunk_index, doc_type, section_title và ngày ban hành. Metadata này không chỉ phục vụ lọc mà còn là nền tảng cho phân vùng đồ thị theo phòng ban.')

# 2.3
add_heading2(doc, '2.3. Xây dựng hệ thống Graph RAG')
add_heading3(doc, '2.3.1. Thiết kế Schema Knowledge Graph')
add_normal(doc, 'Đồ thị Tri thức được thiết kế gồm hai loại thực thể chính: Node văn bản (mỗi chunk là một nút mang đầy đủ content và metadata) và Node thực thể (học phần, điều khoản, phòng ban được nội suy tự động). Hệ thống định nghĩa ba loại cạnh:')
add_bullet(doc, 'Cạnh cấu trúc: Kết nối các chunks tuần tự trong cùng tài liệu (FOLLOWS) hoặc quan hệ bao hàm (PART_OF).')
add_bullet(doc, 'Cạnh quan hệ logic: Kết nối thực thể có liên kết ngữ nghĩa (ví dụ: Học phần A REQUIRES Học phần B).')
add_bullet(doc, 'Cạnh tương đồng ngữ nghĩa: Được tạo khi cosine similarity giữa hai vector nhúng ≥ 0.7 (giới hạn tối đa 5 cạnh mỗi nút để tránh bùng nổ tổ hợp).')
add_normal(doc, 'Sau khi xây dựng đồ thị, thuật toán Louvain được áp dụng để phát hiện cộng đồng (Community Detection) – tự động nhóm các nút kết nối chặt thành các cộng đồng tương ứng với phòng ban. Kỹ thuật Phân vùng đồ thị (Graph Partitioning) tách thành các đồ thị con theo phòng ban để tránh xung đột ngữ nghĩa giữa các quy định cùng tên nhưng khác bộ phận.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 2.2 – Mẫu đồ thị tri thức]', first_indent=False)
add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 2.3 – Đồ thị sau khi phân cụm cộng đồng]', first_indent=False)

add_heading3(doc, '2.3.2. Xây dựng hệ thống Retrieval ba tầng')
add_normal(doc, 'Hệ thống tra cứu vận hành tuần tự qua 3 tầng phân tích kết hợp:')
add_bullet(doc, 'Tầng 1 – Định tuyến Cộng đồng: Câu hỏi được mã hóa thành vector Q, tính độ tương đồng với vector trọng tâm từng cộng đồng để xác định top-3 cộng đồng liên quan nhất, thu hẹp không gian tìm kiếm.')
add_bullet(doc, 'Tầng 2 – Truy xuất Lai (Hybrid Retrieval): Trong top-3 cộng đồng, kết hợp Dense Vector Search (FAISS cosine similarity) và BM25 Sparse Search; điểm tổng hợp = α × vector_score + (1-α) × bm25_score để lấy top-10 tài liệu ứng viên.')
add_bullet(doc, 'Tầng 3 – Duyệt Đồ thị: Từ các tài liệu điểm cao, BFS duyệt các cạnh ngữ nghĩa (cosine ≥ 0.8) để thu thập thêm nút lân cận mang ngữ cảnh liên quan mà tìm kiếm vector thuần không phát hiện được.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 2.4 – Sơ đồ luồng truy vấn Graph RAG]', first_indent=False)

add_heading3(doc, '2.3.3. Tích hợp LLM và xây dựng Agent System')
add_normal(doc, 'Supervisor Agent duy trì trạng thái ngữ cảnh bằng cơ chế Cửa sổ trượt (Sliding Window): khi lịch sử chat quá dài, một LLM phụ tóm lược hội thoại cũ, bảo toàn thông tin cốt lõi trong giới hạn cửa sổ. Khi phát hiện nhân xưng mơ hồ, agent tự giải quyết đồng tham chiếu trước khi định tuyến.')
add_normal(doc, 'ReAct Agent vận hành theo vòng lặp: Suy luận (Thought) – phân rã câu hỏi phức tạp → Hành động (Action) – gọi công cụ (hybrid_search, graph_traversal, score_db) với tham số chính xác → Quan sát (Observation) – đánh giá kết quả và quyết định tiếp tục hay dừng. Kết quả ngữ cảnh cuối cùng được truyền vào LLM để sinh câu trả lời bằng tiếng Việt tự nhiên.')

add_normal(doc, '[CẦN BỔ SUNG THÔNG TIN: Chèn hình Hình 2.5 – Luồng hoạt động của ReAct Agent]', first_indent=False)

add_heading2(doc, '2.4. Tổng kết chương 2')
add_normal(doc, 'Chương 2 đã phác họa và luận giải chi tiết toàn bộ kiến trúc lõi của hệ thống hỏi đáp tư vấn học vụ. Bước đột phá học thuật nằm ở việc ứng dụng Graph RAG thay cho Vector RAG truyền thống, kết hợp Đồ thị Tri thức phân vùng theo phòng ban với cơ chế truy xuất lai ba tầng. Kiến trúc Multi-Agent với ReAct Pattern đảm bảo khả năng xử lý hội thoại đa lượt, tự động định tuyến và kiểm soát chất lượng câu trả lời.')

doc.save('document/TAI_LIEU_SAN_PHAM_CHUAN_FORMAT_part2.docx')
print("Saved part2 – Chuong 1 + Chuong 2 done.")
